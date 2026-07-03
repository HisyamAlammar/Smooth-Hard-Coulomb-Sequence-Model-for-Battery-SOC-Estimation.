from __future__ import annotations

import csv
import json
import math
import shutil
import zipfile
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"C:/Users/VICTUS/Downloads/Template-JNTETI-2025-ENG (3).docx")
OUT = ROOT / "drafts" / "JNTETI_SOC_Hard_Coulomb_v5_Draft_ID.docx"
TRACE = ROOT / "drafts" / "JNTETI_SOC_Hard_Coulomb_v5_Draft_ID_traceability.md"
FIG_OUT = ROOT / "drafts" / "figures_generated"
FIG_SRC = ROOT / "results" / "v5" / "figures"


def read_csv(path: str) -> list[dict[str, str]]:
    with (ROOT / path).open(newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def read_json(path: str):
    with (ROOT / path).open(encoding="utf-8") as f:
        return json.load(f)


def num(x, nd=2) -> str:
    try:
        v = float(x)
        if math.isnan(v):
            return ""
        return f"{v:.{nd}f}"
    except (TypeError, ValueError):
        return str(x)


def style_run(run, font="Times New Roman", size=8.5, bold=False, italic=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_p_style(p, size=8.5, bold=False, italic=False, align=None, font="Times New Roman"):
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(3)
    pf.line_spacing = 1.0
    for r in p.runs:
        style_run(r, font, size, bold, italic)


def set_cell_text(cell, text, bold=False, size=7.2, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    style_run(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "A6A6A6")


def set_columns(section, num_cols=None):
    cols = section._sectPr.xpath("./w:cols")
    if not cols:
        cols_el = OxmlElement("w:cols")
        section._sectPr.append(cols_el)
    else:
        cols_el = cols[0]
    if num_cols is None:
        if qn("w:num") in cols_el.attrib:
            del cols_el.attrib[qn("w:num")]
    else:
        cols_el.set(qn("w:num"), str(num_cols))
    cols_el.set(qn("w:space"), "403")


def clear_doc(doc: Document):
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_par(doc, text="", style="IEEE Paragraph", size=8.5, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(style=style if style in [s.name for s in doc.styles] else None)
    run = p.add_run(text)
    set_p_style(p, size=size, bold=bold, italic=italic, align=align)
    return p


def add_heading(doc, text, level=1):
    style = {1: "IEEE Heading 1", 2: "IEEE Heading 2", 3: "IEEE Heading 3"}[level]
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_p_style(p, size=8.5, bold=True, italic=(level == 2), align=WD_ALIGN_PARAGRAPH.LEFT, font="Helvetica")
    return p


def add_equation(doc, text, number):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text + f"    ({number})")
    style_run(r, font="Cambria Math", size=8.2)
    p.paragraph_format.space_after = Pt(4)


def add_caption(doc, text, kind="Gambar"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    style_run(r, font="Helvetica" if kind == "Gambar" else "Times New Roman", size=7.0, bold=kind == "Gambar")
    p.paragraph_format.space_after = Pt(4)


def add_table(doc, caption, headers, rows, widths=None):
    add_caption(doc, caption, "Tabel")
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=6.8, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], val, size=6.8, align=align)
    return table


def add_figure(doc, path, caption, width=3.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption, "Gambar")


def make_flowchart(path: Path):
    FIG_OUT.mkdir(parents=True, exist_ok=True)
    stages = [
        "Seq2Point\nartifact",
        "Vanilla\nphysics blind",
        "Soft-PINN\nunstable",
        "Post-hoc\nclamp fails",
        "Hard-Coulomb\nsign consistency",
        "Anchor trap\nfound",
        "Dataset v5c\ncorrected",
        "anchor_last\nobservability",
        "Carried\ninference",
        "eta*=2.0\ncalibration",
        "Final:\nanchor_last +\ncalibrated carried",
    ]
    fig, ax = plt.subplots(figsize=(12.0, 3.2), dpi=300)
    ax.axis("off")
    colors = ["#E8EEF7", "#F7E8E8", "#F7F1D8", "#F7E8E8", "#E3F2EA", "#F2E8F7",
              "#E8F1F7", "#DFF0EA", "#EAF0DF", "#FFF0D9", "#D8EAD8"]
    xs = [0.04 + i * 0.088 for i in range(len(stages))]
    y = 0.54
    for i, (x, label) in enumerate(zip(xs, stages)):
        ax.text(x, y, label, ha="center", va="center", fontsize=8.5,
                bbox=dict(boxstyle="round,pad=0.28", fc=colors[i], ec="#4A5568", lw=0.8))
        if i < len(stages) - 1:
            ax.annotate("", xy=(xs[i + 1] - 0.038, y), xytext=(x + 0.038, y),
                        arrowprops=dict(arrowstyle="->", lw=1.0, color="#4A5568"))
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def make_architecture(path: Path):
    fig, ax = plt.subplots(figsize=(8.4, 4.6), dpi=300)
    ax.axis("off")
    boxes = {
        "Input window\nV, I, T, features": (0.06, 0.68, 0.19, 0.14, "#E8EEF7"),
        "2-layer LSTM\nhidden sequence h_1:T": (0.32, 0.68, 0.22, 0.14, "#E3F2EA"),
        "Delta head\nmagnitude logits": (0.62, 0.78, 0.22, 0.12, "#FFF0D9"),
        "Anchor head\nh_T (anchor_last)": (0.62, 0.56, 0.22, 0.12, "#F2E8F7"),
        "Hard-Coulomb route\nsign(I), |I| eta gamma": (0.34, 0.34, 0.26, 0.14, "#E8F1F7"),
        "Cumulative SOC\ns_hat_t": (0.70, 0.34, 0.20, 0.14, "#D8EAD8"),
        "Carried recursive state\n+ eta*=2.0 calibration": (0.36, 0.10, 0.32, 0.13, "#F7F1D8"),
    }
    for label, (x, y, w, h, c) in boxes.items():
        ax.add_patch(plt.Rectangle((x, y), w, h, facecolor=c, edgecolor="#334155", linewidth=0.9))
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=9)
    arrows = [
        ((0.25, 0.75), (0.32, 0.75)),
        ((0.54, 0.75), (0.62, 0.84)),
        ((0.54, 0.75), (0.62, 0.62)),
        ((0.73, 0.78), (0.50, 0.48)),
        ((0.73, 0.56), (0.78, 0.48)),
        ((0.60, 0.41), (0.70, 0.41)),
        ((0.78, 0.34), (0.58, 0.23)),
        ((0.48, 0.23), (0.48, 0.34)),
    ]
    for a, b in arrows:
        ax.annotate("", xy=b, xytext=a, arrowprops=dict(arrowstyle="->", lw=1.1, color="#475569"))
    ax.text(0.06, 0.24, "Klaim: konsistensi tanda terhadap arus terukur,\nbukan kebenaran fisik penuh saat sensor gagal.",
            fontsize=8.5, ha="left", va="center", color="#334155")
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def build():
    FIG_OUT.mkdir(parents=True, exist_ok=True)
    flow = FIG_OUT / "fig_research_evolution_flowchart.png"
    arch = FIG_OUT / "fig_final_architecture.png"
    make_flowchart(flow)
    make_architecture(arch)

    multiseed = read_csv("results/v5/multiseed/multiseed_summary.csv")
    final = read_csv("results/v5/final_v5_model_comparison.csv")
    recursive = read_csv("results/v5/recursive_inference/recursive_policy_comparison.csv")
    eta_rows = read_csv("results/v5/delta_calibration/eta_gamma_sweep.csv")
    ekf_rows = read_csv("results/v5/ekf_ecm/recursive_vs_ekf_comparison.csv")
    dataset_rows = read_csv("results/v5/dataset_variant_comparison.csv")
    claims = read_json("reports/v5_campaign/claims_register_v2.json")
    ranking = read_json("results/v5/multiseed/ranking_stability.json")

    def final_row(model, scenario=None, policy=None):
        for r in final:
            if r["model"] == model and (scenario is None or r["scenario"] == scenario) and (policy is None or r["inference_policy"] == policy):
                return r
        return {}

    def ms(model, scenario):
        for r in multiseed:
            if r["model"] == model and r["scenario"] == scenario:
                return r
        return {}

    anchor_a, anchor_b = ms("hc_anchor_last", "scenario_A"), ms("hc_anchor_last", "scenario_B")
    original_b = ms("hard_coulomb_lstm", "scenario_B")
    eta2 = [r for r in eta_rows if r["eta"] == "2.0" and r["gamma_mode"] == "nominal" and r["mode"] == "inference_sweep"][0]
    load_gated = [r for r in recursive if r["policy"] == "load_gated"][0]
    best_ekf = [r for r in ekf_rows if r["model"] == "EKF_1RC_cont[R=0.01]"][0]

    doc = Document(TEMPLATE)
    clear_doc(doc)
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.3)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(1.3)
    sec.right_margin = Cm(1.3)
    set_columns(sec, None)

    # Front matter
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Estimasi SOC Anchor-Aware Hard-Coulomb untuk Baterai Li-Ion")
    style_run(r, font="Helvetica", size=20, bold=True)
    p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Abyan Hisyam Al'ammar")
    style_run(r, font="Helvetica", size=9)
    add_par(doc, "Program Studi Informatika, Universitas AMIKOM Yogyakarta, Yogyakarta, Indonesia", size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_par(doc, "[Received: DD MM YY, Revised: DD MM YY, Accepted: DD MM YY]", size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_par(doc, "Corresponding Author: Abyan Hisyam Al'ammar (email: [ISI EMAIL])", size=8, align=WD_ALIGN_PARAGRAPH.LEFT)

    abstract = (
        "ABSTRAK - Estimasi state of charge (SOC) yang akurat dan konsisten merupakan komponen penting pada battery "
        "management system baterai litium-ion, terutama saat sel beroperasi pada suhu ekstrem. Model neural berbasis "
        "jendela sering mencapai galat rata-rata yang kompetitif, tetapi dapat menghasilkan lintasan SOC yang tidak "
        "konsisten dengan arah arus terukur dan mengalami cold-start anchor berulang pada protokol Seq2Point. Penelitian "
        "ini menyajikan evolusi lengkap dari baseline LSTM/TCN, soft-constraint, post-hoc clamp, Hard-Coulomb awal, "
        "hingga sistem akhir `anchor_last + calibrated carried inference`. Dataset v5c digunakan untuk mengoreksi "
        "kontaminasi label akibat loaded-start ohmic bias dan konflik decimation. Hard-Coulomb menegakkan konsistensi "
        "tanda delta SOC terhadap arus terukur melalui struktur forward pass, sedangkan anchor_last membaca representasi "
        "jendela penuh untuk meningkatkan observabilitas anchor. Pada evaluasi windowed multi-seed, anchor_last mencapai "
        f"RMSE Scenario A {num(anchor_a['rmse_pct_mean'],3)} +/- {num(anchor_a['rmse_pct_std'],3)}% dan Scenario B "
        f"{num(anchor_b['rmse_pct_mean'],3)} +/- {num(anchor_b['rmse_pct_std'],3)}%. Inference carried yang dikalibrasi "
        f"dengan eta*=2.0 memulihkan rasio delta dari 0,751 menjadi {num(eta2['gated_delta_ratio'],3)} dan menurunkan RMSE "
        f"recursive menjadi {num(eta2['rec_rmse_pct'],3)}%, termasuk {num(eta2['rec_rmse_n20'],3)}% pada -20 C. Dalam "
        f"protokol v5, sistem ini mengungguli baseline EKF terbaik yang diuji ({num(best_ekf['rmse_pct'],3)}% RMSE), "
        "tetapi klaim dibatasi pada dataset, protokol, dan asumsi arus terukur yang diaudit."
    )
    add_par(doc, abstract, style="IEEE Abtract", size=8, bold=False)
    add_par(doc, "KATA KUNCI - Estimasi SOC, Baterai Li-Ion, Hard-Coulomb, LSTM, Anchor-Aware, Recursive Inference, EKF, Suhu Ekstrem.", style="IEEE Abtract", size=8, bold=False)

    doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(doc.sections[-1], 2)

    add_heading(doc, "I. PENDAHULUAN", 1)
    intro_pars = [
        "SOC tidak dapat diukur langsung sehingga harus diinferensikan dari tegangan, arus, temperatur, dan model internal. Pada BMS kendaraan listrik, galat SOC memengaruhi estimasi jarak tempuh, batas charge/discharge, dan keputusan proteksi. Tantangan meningkat pada suhu rendah karena resistansi internal, polarisasi, dan histeresis tegangan memperlemah hubungan terminal voltage terhadap SOC sebenarnya [1]-[5], [8]-[12].",
        "Model sekuens seperti LSTM, TCN, dan varian CNN-LSTM telah digunakan untuk estimasi SOC karena mampu menangkap dependensi temporal nonlinear [6], [7], [13]-[16], [20]. Namun, model neural tanpa batasan dapat menaikkan SOC ketika arus menunjukkan discharge, atau menurunkan SOC ketika arus menunjukkan charge. RMSE yang rendah tidak otomatis menjamin lintasan SOC yang konsisten.",
        "Protokol windowed atau Seq2Point memperparah masalah lain: setiap jendela memiliki cold-start anchor sendiri. Jika anchor awal salah, terutama ketika jendela dimulai saat sel berada di bawah beban dingin, kesalahan dapat muncul berulang. Karena itu, penelitian ini mengevaluasi bukan hanya arsitektur, tetapi juga kebijakan inference yang membawa state SOC antarjendela.",
        "EKF dan ECM tetap relevan sebagai pembanding karena keduanya adalah keluarga estimator rekursif klasik. EKF memberi koreksi tegangan kontinu, tetapi performanya bergantung pada model OCV/ECM, identifikasi parameter, dan pengaturan noise measurement. V5 membandingkan Hard-Coulomb recursive terhadap EKF kontinu pada rantai evaluasi yang sama.",
    ]
    for t in intro_pars:
        add_par(doc, t)
    add_par(doc, "Kontribusi penelitian ini adalah sebagai berikut.", bold=False)
    contribs = [
        "Estimator neural Hard-Coulomb yang menegakkan konsistensi tanda terhadap arus terukur.",
        "Ablation bertahap yang menunjukkan kegagalan vanilla LSTM/TCN, soft penalty, post-hoc clamp, dan anchor-first.",
        "Evaluasi dataset v5c yang memperbaiki artifact label dan decimation dari v4.",
        "Desain anchor_last dan strategi `anchor_last + calibrated carried inference`.",
        "Pembandingan terhadap null OCV+Coulomb dan baseline EKF/1RC kontinu.",
    ]
    for i, t in enumerate(contribs, 1):
        add_par(doc, f"{i}. {t}", align=WD_ALIGN_PARAGRAPH.LEFT)

    add_heading(doc, "II. METODE", 1)
    add_heading(doc, "A. DATASET DAN KOREKSI V5", 2)
    add_par(doc, "Kampanye v5 dibangun untuk menutup tiga ancaman validitas dari v4: label terkontaminasi saat segment start berada di bawah beban, konflik routing akibat decimation first-sample, dan perbedaan antara evaluasi windowed dan deployment recursive. Varian final v5c memakai label ohmic-corrected dan mean-per-second decimation; jumlah segmen serta window split tetap identik agar perubahan dapat diatribusikan ke koreksi data, bukan split baru.")
    add_par(doc, "Pada laporan Phase 1, 48/103 segmen legacy memiliki loaded-start anchor bias dengan mean 2,4% SOC, p95 10,9%, dan worst 33,5%. Mean-per-second decimation menghilangkan routing-sign conflict dan envelope-unsatisfiable seconds by construction karena arus yang dipakai adalah rerata intradetik.")
    ds_rows = [[r["variant"], r["scenario"], r["label_mode"], r["decimation_mode"], r["segments"], f"{r['train_windows']}/{r['val_windows']}/{r['test_windows']}", num(r["label_shift_vs_v4_mean_pct"], 2)] for r in dataset_rows if r["scenario"] == "A"]
    add_table(doc, "Tabel I. Perbandingan Varian Dataset Scenario A", ["Varian", "Scen", "Label", "Decimation", "Seg", "Tr/Val/Test", "Shift"], ds_rows)

    add_heading(doc, "B. FORMULASI SOC DAN COULOMB COUNTING", 2)
    add_par(doc, "Coulomb counting memperbarui SOC dengan mengintegrasikan arus terhadap waktu dan kapasitas efektif temperatur. Dalam naskah ini arus discharge mengikuti konvensi negatif yang digunakan oleh pipeline evaluasi.")
    add_equation(doc, r"SOC_t = SOC_0 - \frac{\eta \Delta t}{3600 Q_{\mathrm{eff}}(T)} \sum_{k=1}^{t} I_k", 1)
    add_par(doc, "Pada (1), SOC_t adalah state of charge pada timestep t, I_k adalah arus terukur, Delta t adalah interval sampling, Q_eff(T) adalah kapasitas efektif pada temperatur T, dan eta adalah faktor skala laju.")

    add_heading(doc, "C. HARD-COULOMB CONSTRAINT", 2)
    add_par(doc, "Hard-Coulomb memindahkan prior fisika dari loss function ke ruang keluaran model. Model tetap belajar magnitude delta, tetapi tanda delta ditentukan oleh arus terukur. Akibatnya, PVR 0,00% adalah sifat struktural forward pass, bukan keberhasilan empiris yang boleh diklaim sebagai hasil akurasi.")
    add_equation(doc, r"m_t^{dis} = \mathbb{1}[I_t < -I_{\epsilon}], \quad m_t^{chg} = \mathbb{1}[I_t > I_{\epsilon}]", 2)
    add_equation(doc, r"\Delta \hat{s}_t = \begin{cases} -|\Delta \hat{s}_t|, & I_t < -I_{\epsilon} \\ |\Delta \hat{s}_t|, & I_t > I_{\epsilon} \\ 0\ \mathrm{atau\ bounded\ drift}, & |I_t| \le I_{\epsilon} \end{cases}", 3)
    add_equation(doc, r"\hat{s}_t = \hat{s}_0 + \sum_{k=1}^{t}\Delta \hat{s}_k", 4)
    add_par(doc, "Jaminan ini terbatas pada konsistensi tanda terhadap arus yang diukur. Jika sensor arus salah, offset, atau fault, jaminan measured-current consistency tidak identik dengan kebenaran fisik sel.")

    add_heading(doc, "D. ANCHOR-AWARE MODEL", 2)
    add_par(doc, "Model HC awal memakai anchor head pada h_1. Source code menunjukkan anchor_logit = anchor_head(h[:, 0, :]); akibatnya anchor hanya melihat konteks awal jendela. Varian anchor_last memakai h_T sehingga anchor dikondisikan pada seluruh jendela kausal sebelum diremap ke interval feasible.")
    add_equation(doc, r"\hat{s}_0 = f_{\theta}(h_T) = f_{\theta}(\mathrm{Encoder}(x_{1:T}))", 5)

    add_heading(doc, "E. CALIBRATED CARRIED INFERENCE", 2)
    add_par(doc, "Window independent inference mengulang cold-start pada setiap window. Carried inference membawa state SOC antarwindow kontinu, sehingga hanya awal rantai yang memakai anchor terpelajari. Load-gated policy menjadi fallback ketika kontinuitas rantai tidak pasti, dengan re-anchor pada kondisi rest atau low-load.")
    add_equation(doc, r"\hat{s}^{rec}_{t} = \hat{s}^{rec}_{t-1} + \eta^{*}\Delta \hat{s}_t", 6)
    add_equation(doc, r"r_{\Delta} = \frac{\mathbb{E}(|\Delta \hat{s}_t|)}{\mathbb{E}(|\Delta s_t|)}", 7)
    add_par(doc, f"Pada checkpoint Scenario A seed 42, eta*=2,0 mengubah delta ratio dari 0,751 menjadi {num(eta2['gated_delta_ratio'],3)} dan recursive RMSE menjadi {num(eta2['rec_rmse_pct'],3)}%. Bukti ini dilabeli single-checkpoint/single-scenario sampai re-derivation validasi dan konfirmasi multi-seed selesai.")

    add_heading(doc, "F. BASELINE METHODS", 2)
    add_par(doc, "Baseline meliputi null OCV+Coulomb, Vanilla LSTM, Vanilla TCN, post-hoc clamp pada keluaran vanilla, HC-LSTM anchor-first, HC-TCN, varian anchor_pooled, dan EKF kontinu berbasis OCV-Rint serta 1RC-ECM. EKF memakai OCV 25 C, tabel Q(T), tau 50 s, R1=0,5 R_int(T), dan sweep measurement-noise R dalam {1e-4, 9e-4, 1e-2}.")
    add_figure(doc, arch, "Gambar 1. Arsitektur sistem akhir anchor_last + calibrated carried inference.", width=3.1)

    add_heading(doc, "III. SETUP EKSPERIMEN", 1)
    add_par(doc, "Scenario A adalah temperature-OOD dengan evaluasi pada suhu yang ditahan, termasuk -20 C. Scenario B adalah split in-distribution kronologis. Model utama dievaluasi pada v5c dengan seeds 1-5. Eksperimen recursive, eta calibration, dan EKF menggunakan Scenario A seed 42 sebagai kampanye deployment kontinu tanpa retraining berat.")
    add_equation(doc, r"RMSE = \sqrt{\frac{1}{N}\sum_{i=1}^{N}(s_i-\hat{s}_i)^2}", 8)
    add_equation(doc, r"MAE = \frac{1}{N}\sum_{i=1}^{N}|s_i-\hat{s}_i|", 9)
    add_equation(doc, r"MaxE = \max_i |s_i-\hat{s}_i|", 10)
    add_equation(doc, r"PVR_{dis} = \frac{\sum_t \mathbb{1}[I_t < -I_{\epsilon}] \mathbb{1}[\Delta \hat{s}_t > \epsilon]}{\sum_t \mathbb{1}[I_t < -I_{\epsilon}]}", 11)
    add_equation(doc, r"r_{\Delta} = \frac{\mathrm{mean}(|\Delta \hat{s}|)}{\mathrm{mean}(|\Delta s|)}", 12)

    add_heading(doc, "IV. HASIL DAN PEMBAHASAN", 1)
    add_heading(doc, "A. EVOLUSI PENELITIAN", 2)
    add_figure(doc, flow, "Gambar 2. Alur evolusi penelitian dari artifact Seq2Point menuju sistem akhir.", width=3.2)
    add_par(doc, "Naskah ini sengaja menyajikan kegagalan bertahap karena sistem akhir tidak ditemukan melalui pemilihan model arbitrer. Seq2Point/windowed protocol memunculkan artifact cold-start berulang. Vanilla LSTM/TCN memperlihatkan physics blindness. Soft-PINN dan penalty-based constraint tidak cukup stabil karena penalty hanya preferensi optimisasi. Post-hoc clamp memperbaiki tanda setelah prediksi, tetapi merusak lintasan dan menghasilkan RMSE besar. Hard-Coulomb kemudian memberi sign consistency by construction, tetapi versi awal anchor-first terbatas oleh observabilitas anchor.")
    phase_rows = [
        ["P0", "Legacy freeze", "v4 preserved"],
        ["P1", "Dataset v5c", "label/decimation fixed"],
        ["P2", "Headline retrain", "v5 models"],
        ["P3", "Multi-seed", "seeds 1-5"],
        ["P4", "Recursive policies", "load_gated"],
        ["P5", "Eta calibration", "eta*=2.0"],
        ["P6", "EKF baselines", "continuous EKF"],
        ["P7", "Final matrix", "62 rows"],
        ["P8", "Claims register", "17 claims"],
        ["P9", "Final report", "rewrite brief"],
        ["P10", "Readiness gate", "52/52 PASS"],
    ]
    add_table(doc, "Tabel II. Checklist Fase Kampanye v5", ["Fase", "Artefak", "Status"], phase_rows)

    add_heading(doc, "B. PENGARUH KOREKSI DATASET V5", 2)
    add_figure(doc, FIG_SRC / "soc_initial_bias_by_temperature.png", "Gambar 3. Bias initial SOC pada segment start menurut temperatur.", width=3.0)
    add_figure(doc, FIG_SRC / "routing_conflict_by_decimation_mode.png", "Gambar 4. Konflik routing akibat decimation legacy dan koreksi mean-per-second.", width=3.0)
    add_par(doc, "Koreksi v5c menurunkan sebagian narasi catastrophic MaxE v4 yang terbukti artifact label. Namun, kesimpulan utama tidak dibalik: HC anchor variants tetap lebih baik daripada null dan vanilla pada matriks v5c, sedangkan original HC Scenario B justru terkonfirmasi gagal sistematis.")

    add_heading(doc, "C. POST-HOC CLAMP DAN TRAINING-THROUGH-CONSTRAINT", 2)
    clamp_a = final_row("vanilla+posthoc_clamp", "scenario_A")
    clamp_b = final_row("vanilla+posthoc_clamp", "scenario_B")
    add_par(doc, f"Post-hoc clamp gagal pada label terkoreksi: RMSE Scenario A {num(clamp_a['rmse_pct'],2)}% dan Scenario B {num(clamp_b['rmse_pct'],2)}%. Nilai ini jauh lebih buruk daripada HC windowed. Hasil ini mendukung klaim bahwa constraint perlu dilalui saat training/forward computation, bukan ditempel sebagai filter setelah model bebas menghasilkan lintasan.")

    add_heading(doc, "D. STABILITAS MULTI-SEED DAN ANCHOR-LAST", 2)
    ms_rows = []
    names = ["vanilla_lstm", "hard_coulomb_lstm", "hard_coulomb_tcn", "hc_anchor_pooled", "hc_anchor_last"]
    for name in names:
        a, b = ms(name, "scenario_A"), ms(name, "scenario_B")
        ms_rows.append([name, f"{num(a['rmse_pct_mean'])} +/- {num(a['rmse_pct_std'])}", f"{num(a['maxe_pct_mean'])} +/- {num(a['maxe_pct_std'])}", f"{num(b['rmse_pct_mean'])} +/- {num(b['rmse_pct_std'])}", f"{num(b['maxe_pct_mean'])} +/- {num(b['maxe_pct_std'])}"])
    cmp_models = ["null[ocv25_qnom]", "vanilla_lstm", "hard_coulomb_lstm", "hc_anchor_last", "vanilla+posthoc_clamp"]
    cmp_rows = []
    for model_name in cmp_models:
        a, b = final_row(model_name, "scenario_A"), final_row(model_name, "scenario_B")
        cmp_rows.append([model_name, num(a.get("rmse_pct", "")), num(a.get("maxe_pct", "")), num(b.get("rmse_pct", "")), num(b.get("maxe_pct", ""))])
    add_table(doc, "Tabel III. Perbandingan Model dan Baseline Windowed", ["Model", "A RMSE", "A MaxE", "B RMSE", "B MaxE"], cmp_rows)
    add_table(doc, "Tabel IV. Ringkasan Multi-Seed v5c", ["Model", "A RMSE", "A MaxE", "B RMSE", "B MaxE"], ms_rows)
    add_figure(doc, FIG_SRC / "multiseed_rmse_boxplot.png", "Gambar 5. Distribusi RMSE multi-seed untuk model utama.", width=3.2)
    add_par(doc, f"Anchor_last mencapai Scenario A {num(anchor_a['rmse_pct_mean'],3)} +/- {num(anchor_a['rmse_pct_std'],3)}% dan Scenario B {num(anchor_b['rmse_pct_mean'],3)} +/- {num(anchor_b['rmse_pct_std'],3)}%. Original HC-LSTM gagal sistematis pada Scenario B dengan {num(original_b['rmse_pct_mean'],3)} +/- {num(original_b['rmse_pct_std'],3)}%, dan laporan Phase 3 mencatat kegagalan pada 5/5 seeds. Ranking seed menunjukkan anchor_last menang 4/5 seed Scenario A; pada Scenario B, keluarga anchor_last/pooled menang 5/5 seed.")

    add_heading(doc, "E. RECURSIVE INFERENCE", 2)
    rec_rows = [[r["policy"], num(r["rmse_pct"]), num(r["maxe_pct"]), num(r["rmse_n20degC"]), num(r["rmse_40degC"]), num(r["reanchor_pct"])] for r in recursive]
    add_table(doc, "Tabel V. Perbandingan Kebijakan Recursive Scenario A Seed 42", ["Policy", "RMSE", "MaxE", "-20 C", "40 C", "Re-anchor"], rec_rows)
    add_figure(doc, FIG_SRC / "recursive_policy_temperature_breakdown.png", "Gambar 6. Breakdown temperatur untuk kebijakan recursive.", width=3.2)
    add_par(doc, f"Load-gated policy adalah fixed gate terbaik sebelum eta calibration, dengan RMSE {num(load_gated['rmse_pct'],3)}% atau sekitar -2,64 pp terhadap windowed independent. Pure carried anchor mengurangi error -20 C, tetapi memperburuk drift hangat; ini menunjukkan bahwa state carry harus dipasangkan dengan kalibrasi delta path atau gate fisik.")

    add_heading(doc, "F. KALIBRASI ETA DAN DELTA-PATH", 2)
    eta_show = [r for r in eta_rows if r["mode"] == "inference_sweep" and r["gamma_mode"] == "nominal" and r["eta"] in {"1.287", "1.5", "1.75", "2.0", "2.5", "3.0"}]
    eta_table = [[r["eta"], num(r["gated_delta_ratio"], 3), num(read_json("results/v5/delta_calibration/eta_gamma_sweep.json")["rows"][eta_rows.index(r)]["rec_rmse_pct"], 2), num(read_json("results/v5/delta_calibration/eta_gamma_sweep.json")["rows"][eta_rows.index(r)]["rec_rmse_n20"], 2), num(read_json("results/v5/delta_calibration/eta_gamma_sweep.json")["rows"][eta_rows.index(r)]["rec_rmse_40"], 2)] for r in eta_show]
    add_table(doc, "Tabel VI. Sweep Eta Nominal pada Inference-Time Calibration", ["Eta", "Delta ratio", "Rec RMSE", "-20 C", "40 C"], eta_table)
    add_figure(doc, FIG_SRC / "eta_vs_delta_ratio.png", "Gambar 7. Eta calibration memulihkan rasio delta-path.", width=3.0)
    add_figure(doc, FIG_SRC / "eta_vs_rmse_by_temperature.png", "Gambar 8. Pengaruh eta terhadap RMSE recursive per temperatur.", width=3.0)
    add_par(doc, f"Eta*=2,0 nominal merupakan titik rate fidelity: rasio delta mendekati 1 dan recursive RMSE mencapai {num(eta2['rec_rmse_pct'],3)}%. Retraining pada eta=2,0 tidak otomatis self-calibrate karena magnitude head mengompensasi envelope; laporan Phase 5 menunjukkan recursive retrained eta=2,0 tetap sekitar 14,91%, jauh di atas inference-time calibration.")

    add_heading(doc, "G. PERBANDINGAN DENGAN EKF", 2)
    ekf_table = [[r["model"], num(r["rmse_pct"]), num(r["maxe_pct"]), num(r["rmse_n20"]), num(r["rmse_40"]), num(r["pvr_disch_eps0"])] for r in ekf_rows]
    ekf_table.insert(3, ["HC carried @ eta*=2.0", num(eta2["rec_rmse_pct"]), num(eta2["rec_maxe_pct"]), num(eta2["rec_rmse_n20"]), num(eta2["rec_rmse_40"]), "0.00"])
    add_table(doc, "Tabel VII. Hard-Coulomb Recursive dan EKF Kontinu", ["Model", "RMSE", "MaxE", "-20 C", "40 C", "PVR"], ekf_table)
    add_figure(doc, FIG_SRC / "recursive_vs_ekf_temperature_breakdown.png", "Gambar 9. Perbandingan temperatur antara HC recursive dan EKF.", width=3.2)
    add_par(doc, f"EKF terbaik adalah 1RC dengan R=1e-2, RMSE {num(best_ekf['rmse_pct'],3)}%, -20 C {num(best_ekf['rmse_n20'],3)}%, dan PVR discharge {num(best_ekf['pvr_disch_eps0'],3)}%. Pada R lebih kecil, EKF mempercayai tegangan yang terkontaminasi polarisasi dan RMSE naik hingga 39,12%. Karena itu, perbandingan yang aman bukan superioritas universal, melainkan trade-off koreksi tegangan versus konsistensi tanda. Dalam protokol v5 yang diuji, calibrated recursive HC lebih rendah RMSE dan tetap PVR 0,00% by construction.")

    add_heading(doc, "H. ANALISIS KLAIM DAN KETERBATASAN", 2)
    claim_rows = [[c["id"], c["status"], c["claim"][:72]] for c in claims["claims"] if c["id"] in (1, 2, 5, 6, 11, 12, 15, 16, 17)]
    add_table(doc, "Tabel VIII. Ringkasan Klaim Kunci dari Claims Register v2", ["ID", "Status", "Klaim"], claim_rows)
    add_par(doc, "Batasan utama: PVR 0,00% adalah sifat struktural terhadap arus terukur, bukan bukti kebenaran fisik total; sensor fault tidak dijamin; eta* perlu re-derivation pada validation chains dan konfirmasi Scenario B/multi-seed; EKF memakai parameter literature-like, bukan identifikasi cell-specific; deployment edge/TinyML tidak diklaim karena WCET/RAM hardware belum diuji.")

    add_heading(doc, "V. KESIMPULAN", 1)
    add_par(doc, "Penelitian ini menyimpulkan bahwa sistem akhir `anchor_last + calibrated carried inference` adalah hasil dari staged failure analysis dan ablation, bukan pemilihan model arbitrer. Hard-Coulomb menegakkan konsistensi tanda terhadap arus terukur; anchor_last memperbaiki observabilitas anchor; recursive inference mengurangi artifact cold-start; dan eta calibration memperbaiki underestimation delta-path tanpa retraining. Dalam protokol v5, calibrated recursive HC mencapai RMSE 4,4318% dan mengungguli baseline EKF terbaik yang diuji sebesar 6,8485%, dengan caveat bahwa bukti eta masih single-checkpoint/single-scenario.")
    add_par(doc, "Pekerjaan berikutnya perlu mengestimasi eta* dari validation chains, menguji Scenario B dan multi-seed recursive calibration, menambahkan diagnostik sensor fault, menguji quantized deployment dengan sign-preserving accumulator, serta memperluas evaluasi ke dataset dan kimia baterai lain.")

    for title, text in [
        ("KONFLIK KEPENTINGAN", "Penulis menyatakan tidak terdapat konflik kepentingan."),
        ("KONTRIBUSI PENULIS", "Abyan Hisyam Al'ammar: konseptualisasi, metodologi, perangkat lunak, validasi, analisis formal, investigasi, kurasi data, visualisasi, penulisan draf awal, dan penyuntingan naskah. [ISI NAMA DOSEN PEMBIMBING]: supervisi, validasi metodologi, dan penelaahan naskah."),
        ("UCAPAN TERIMA KASIH", "Penulis mengucapkan terima kasih kepada [ISI NAMA DOSEN/PROGRAM STUDI/LAB] atas arahan dan masukan selama proses penelitian."),
    ]:
        add_heading(doc, title, 1)
        add_par(doc, text)

    add_heading(doc, "REFERENSI", 1)
    refs_text = (ROOT / "JNTETI_Hard_Coulomb_LSTM_Draft_ID.md").read_text(encoding="utf-8")
    refs = [line.strip() for line in refs_text.splitlines() if line.strip().startswith("[") and "]" in line[:5]]
    for ref in refs[:24]:
        p = doc.add_paragraph(style="IEEE Reference Item")
        r = p.add_run(ref)
        style_run(r, size=7.5)
        p.paragraph_format.space_after = Pt(2)
    add_par(doc, "[REF NEEDED] Tambahkan referensi terbaru 2024-2026 untuk SOC estimation suhu ekstrem bila data bibliografis terverifikasi tersedia.", style="IEEE Reference Item", size=7.5, align=WD_ALIGN_PARAGRAPH.LEFT)

    add_heading(doc, "LAMPIRAN", 1)
    add_par(doc, f"Final ablation matrix berisi 62 baris; claims register v2 berisi {len(claims['claims'])} klaim; readiness gate Phase 10 melaporkan 52/52 PASS. Notebook bukti final berada di notebooks/ablation_studies_v5_final/12-19.")
    add_par(doc, "Artefak bernama final_v5_research_report.md, manuscript_rewrite_brief.md, manuscript_readiness_gate.md, results/v5/final_figures/, dan results/v5/final_tables/ tidak ditemukan persis sesuai prompt; substitusi aktual dicatat dalam traceability file. [PERLU DIVERIFIKASI]")

    doc.save(OUT)

    # Structural validation.
    Document(OUT)
    with zipfile.ZipFile(OUT) as zf:
        names = set(zf.namelist())
        assert "word/document.xml" in names

    trace_lines = [
        "# Traceability - JNTETI_SOC_Hard_Coulomb_v5_Draft_ID",
        "",
        "## Result Files Used",
        "- reports/v5_campaign/phase1_dataset_v5_report.md",
        "- reports/v5_campaign/phase3_multiseed_report.md",
        "- reports/v5_campaign/phase4_recursive_policies_report.md",
        "- reports/v5_campaign/phase5_delta_calibration_report.md",
        "- reports/v5_campaign/phase6_ekf_baselines_report.md",
        "- reports/v5_campaign/phase7_final_ablation_model_card.md",
        "- reports/v5_campaign/phase9_final_v5_report.md",
        "- reports/v5_campaign/phase10_manuscript_readiness_gate.md",
        "- reports/v5_campaign/claims_register_v2.md/json",
        "- results/v5/dataset_variant_comparison.csv",
        "- results/v5/final_v5_model_comparison.csv",
        "- results/v5/final_ablation_matrix.csv",
        "- results/v5/multiseed/multiseed_summary.csv",
        "- results/v5/multiseed/ranking_stability.json",
        "- results/v5/recursive_inference/recursive_policy_comparison.csv",
        "- results/v5/delta_calibration/eta_gamma_sweep.csv/json",
        "- results/v5/ekf_ecm/recursive_vs_ekf_comparison.csv",
        "- results/v5/ekf_ecm/continuous_ekf_results.json",
        "",
        "## Figures Used",
        f"- {flow.relative_to(ROOT)} (generated)",
        f"- {arch.relative_to(ROOT)} (generated)",
        "- results/v5/figures/soc_initial_bias_by_temperature.png",
        "- results/v5/figures/routing_conflict_by_decimation_mode.png",
        "- results/v5/figures/multiseed_rmse_boxplot.png",
        "- results/v5/figures/recursive_policy_temperature_breakdown.png",
        "- results/v5/figures/eta_vs_delta_ratio.png",
        "- results/v5/figures/eta_vs_rmse_by_temperature.png",
        "- results/v5/figures/recursive_vs_ekf_temperature_breakdown.png",
        "",
        "## Tables Used",
        "- Dataset variants from dataset_variant_comparison.csv",
        "- Phase checklist from v5 campaign reports and readiness gate",
        "- Windowed model/baseline comparison from final_v5_model_comparison.csv",
        "- Multi-seed model comparison from multiseed_summary.csv",
        "- Recursive policy comparison from recursive_policy_comparison.csv",
        "- Eta calibration table from eta_gamma_sweep.csv/json",
        "- EKF comparison from recursive_vs_ekf_comparison.csv",
        "- Claim summary from claims_register_v2.json",
        "",
        "## Missing Artifacts / Aliases",
        "- reports/v5_campaign/final_v5_research_report.md -> substituted with reports/v5_campaign/phase9_final_v5_report.md [PERLU DIVERIFIKASI]",
        "- reports/v5_campaign/manuscript_rewrite_brief.md -> no exact file; Phase 9 contains rewrite brief section [PERLU DIVERIFIKASI]",
        "- reports/v5_campaign/manuscript_readiness_gate.md -> substituted with phase10_manuscript_readiness_gate.md [PERLU DIVERIFIKASI]",
        "- results/v5/final_figures/ -> figures are under results/v5/figures/ [PERLU DIVERIFIKASI]",
        "- results/v5/final_tables/ -> tables are CSV/JSON files under results/v5/ [PERLU DIVERIFIKASI]",
        "",
        "## Manually Inserted Values",
        f"- anchor_last Scenario A {num(anchor_a['rmse_pct_mean'],3)} +/- {num(anchor_a['rmse_pct_std'],3)} from multiseed_summary.csv",
        f"- anchor_last Scenario B {num(anchor_b['rmse_pct_mean'],3)} +/- {num(anchor_b['rmse_pct_std'],3)} from multiseed_summary.csv",
        f"- original HC Scenario B {num(original_b['rmse_pct_mean'],3)} +/- {num(original_b['rmse_pct_std'],3)} from multiseed_summary.csv",
        f"- load_gated RMSE {num(load_gated['rmse_pct'],3)} from recursive_policy_comparison.csv",
        f"- eta*=2.0 rec RMSE {num(eta2['rec_rmse_pct'],4)}, -20 C {num(eta2['rec_rmse_n20'],4)}, delta ratio {num(eta2['gated_delta_ratio'],4)} from eta_gamma_sweep",
        f"- best EKF RMSE {num(best_ekf['rmse_pct'],4)}, PVR {num(best_ekf['pvr_disch_eps0'],4)} from recursive_vs_ekf_comparison.csv",
        "- Author/email/supervisor/acknowledgment placeholders intentionally remain as [ISI ...].",
        "- One reference placeholder intentionally remains as [REF NEEDED].",
        "",
        "## [PERLU DIVERIFIKASI] Notes",
        "- Exact missing path aliases listed above.",
        "- Bibliography was copied from existing repo draft and should be verified before submission.",
        "- Eta* validation-chain re-derivation plus Scenario B/multi-seed confirmation remain open items in readiness gate.",
    ]
    TRACE.write_text("\n".join(trace_lines) + "\n", encoding="utf-8")

    print(OUT)
    print(TRACE)


if __name__ == "__main__":
    build()
