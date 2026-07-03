from __future__ import annotations

import csv
import json
import math
import re
import zipfile
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"C:/Users/VICTUS/Downloads/Template-JNTETI-2025-ENG (3).docx")
OUT = ROOT / "drafts" / "JNTETI_SOC_Hard_Coulomb_v5_Draft_ID_v2.docx"
TRACE = ROOT / "drafts" / "JNTETI_SOC_Hard_Coulomb_v5_Draft_ID_v2_traceability.md"
NOTES = ROOT / "drafts" / "JNTETI_SOC_Hard_Coulomb_v5_Draft_ID_v2_revision_notes.md"
FIG_DIR = ROOT / "drafts" / "figures_v2"


def read_csv(rel: str) -> list[dict[str, str]]:
    with (ROOT / rel).open(newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def read_json(rel: str):
    with (ROOT / rel).open(encoding="utf-8") as f:
        return json.load(f)


def fmt(x, nd=2) -> str:
    try:
        v = float(x)
        if math.isnan(v):
            return ""
        return f"{v:.{nd}f}"
    except (TypeError, ValueError):
        return str(x)


def metric(mean, std=None, nd=2) -> str:
    if std in (None, ""):
        return fmt(mean, nd)
    return f"{fmt(mean, nd)} ± {fmt(std, nd)}"


def run_font(run, font="Times New Roman", size=8.5, bold=False, italic=False):
    run.font.name = font
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def p_format(p, size=8.5, bold=False, italic=False, align=None, font="Times New Roman"):
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    for run in p.runs:
        run_font(run, font, size, bold, italic)


def clear_doc(doc: Document):
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_columns(section, n=None):
    cols = section._sectPr.xpath("./w:cols")
    cols_el = cols[0] if cols else OxmlElement("w:cols")
    if not cols:
        section._sectPr.append(cols_el)
    if n is None:
        cols_el.attrib.pop(qn("w:num"), None)
    else:
        cols_el.set(qn("w:num"), str(n))
    cols_el.set(qn("w:space"), "403")


def add_par(doc, text="", style="IEEE Paragraph", size=8.5, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(style=style if style in [s.name for s in doc.styles] else None)
    r = p.add_run(text)
    p_format(p, size=size, bold=bold, italic=italic, align=align)
    return p


def add_heading(doc, text, level=1):
    style = {1: "IEEE Heading 1", 2: "IEEE Heading 2", 3: "IEEE Heading 3"}[level]
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    p_format(p, size=8.5, bold=True, italic=(level == 2), align=WD_ALIGN_PARAGRAPH.LEFT, font="Helvetica")
    return p


def add_eq(doc, expr, number):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{expr}    ({number})")
    run_font(r, font="Cambria Math", size=8.2)
    p.paragraph_format.space_after = Pt(4)
    return p


def shade(cell, fill="DCEAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False, size=6.6, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text))
    run_font(r, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def borders(table):
    tbl_pr = table._tbl.tblPr
    tb = tbl_pr.first_child_found_in("w:tblBorders")
    if tb is None:
        tb = OxmlElement("w:tblBorders")
        tbl_pr.append(tb)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tb.find(qn("w:" + edge))
        if el is None:
            el = OxmlElement("w:" + edge)
            tb.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "9CA3AF")


def table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    run_font(r, font="Times New Roman", size=7.2, bold=True)
    p.paragraph_format.space_after = Pt(2)


def figure_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    run_font(r, font="Helvetica", size=7.0, bold=True)
    p.paragraph_format.space_after = Pt(4)


def add_table(doc, caption, headers, rows, size=6.4):
    table_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    borders(table)
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, bold=True, size=size, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(table.rows[0].cells[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            cell_text(cells[i], val, size=size, align=align)
    return table


def add_figure(doc, path, caption, width=3.15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    figure_caption(doc, caption)


def save_fig(path):
    plt.savefig(path, dpi=300, bbox_inches="tight")
    plt.close()


def make_figures(data):
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    colors = {
        "ours": "#E76F51",
        "hc": "#2A9D8F",
        "base": "#8A9BA8",
        "ekf": "#6D597A",
        "null": "#B0BEC5",
        "warn": "#E9C46A",
    }

    # Figure 1: architecture diagram
    fig, ax = plt.subplots(figsize=(8.8, 4.8))
    ax.axis("off")
    boxes = [
        ("Jendela input\nV, I, T, fitur", 0.05, 0.70, 0.20, 0.14, "#E8EEF7"),
        ("Encoder LSTM\nh₁ … h_T", 0.32, 0.70, 0.18, 0.14, "#E3F2EA"),
        ("Anchor head\nh_T", 0.62, 0.72, 0.18, 0.12, "#F2E8F7"),
        ("Delta head\n|Δŝ_t|", 0.62, 0.52, 0.18, 0.12, "#FFF0D9"),
        ("Hard-Coulomb\nsign(I_t), |I_t|ηγ", 0.33, 0.38, 0.26, 0.14, "#E8F1F7"),
        ("ŝ_t = anchor + ΣΔŝ", 0.68, 0.34, 0.22, 0.14, "#D8EAD8"),
        ("Carried state\neta*=2.0", 0.37, 0.12, 0.24, 0.12, "#F7F1D8"),
    ]
    for label, x, y, w, h, fill in boxes:
        ax.add_patch(plt.Rectangle((x, y), w, h, facecolor=fill, edgecolor="#334155", lw=0.9))
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=9)
    for a, b in [((0.25, 0.77), (0.32, 0.77)), ((0.50, 0.77), (0.62, 0.78)),
                 ((0.50, 0.73), (0.62, 0.58)), ((0.71, 0.52), (0.56, 0.47)),
                 ((0.80, 0.72), (0.80, 0.48)), ((0.59, 0.45), (0.68, 0.41)),
                 ((0.79, 0.34), (0.55, 0.24)), ((0.49, 0.24), (0.47, 0.38))]:
        ax.annotate("", xy=b, xytext=a, arrowprops=dict(arrowstyle="->", lw=1.1, color="#475569"))
    ax.text(0.05, 0.22, "Jaminan: konsistensi tanda terhadap arus terukur.\nBukan kebenaran fisik penuh saat sensor gagal.",
            fontsize=8.2, color="#334155", ha="left", va="center")
    save_fig(FIG_DIR / "fig1_architecture.png")

    # Figure 2: evolution flowchart
    labels = ["Seq2Point\ncold-start", "Vanilla\nphysics-blind", "Soft penalty\nnon-structural",
              "Post-hoc\nclamp fails", "HC awal\nanchor-first", "v5c\ncorrection",
              "anchor_last", "carried\ninference", "eta*=2.0", "Final system"]
    fig, ax = plt.subplots(figsize=(12, 3.0))
    ax.axis("off")
    xs = np.linspace(0.05, 0.95, len(labels))
    fills = ["#F7E8E8", "#F7E8E8", "#FFF0D9", "#F7E8E8", "#E8F1F7",
             "#E8EEF7", "#D8EAD8", "#E3F2EA", "#FFF0D9", "#D8EAD8"]
    for i, (x, lab) in enumerate(zip(xs, labels)):
        ax.text(x, 0.56, lab, ha="center", va="center", fontsize=8.5,
                bbox=dict(boxstyle="round,pad=0.28", fc=fills[i], ec="#475569", lw=0.8))
        if i < len(labels) - 1:
            ax.annotate("", xy=(xs[i+1] - 0.035, 0.56), xytext=(x + 0.035, 0.56),
                        arrowprops=dict(arrowstyle="->", lw=1.0, color="#475569"))
    save_fig(FIG_DIR / "fig2_evolution.png")

    # Figure 3: dataset correction evidence
    ds = data["dataset"]
    variants = ["v4", "v5a", "v5b", "v5c"]
    vals = [float(next(r for r in ds if r["variant"] == v and r["scenario"] == "A")["label_shift_vs_v4_mean_pct"]) for v in variants]
    fig, ax = plt.subplots(figsize=(4.0, 2.8))
    ax.bar(variants, vals, color=[colors["null"], colors["warn"], "#7FB3D5", colors["ours"]])
    ax.set_ylabel("Mean label shift vs v4 (%SOC)")
    ax.set_xlabel("Dataset variant, Scenario A")
    ax.spines[["top", "right"]].set_visible(False)
    ax.grid(axis="y", alpha=0.2)
    for i, v in enumerate(vals):
        ax.text(i, v + 0.05, f"{v:.2f}", ha="center", fontsize=8)
    save_fig(FIG_DIR / "fig3_dataset_shift.png")

    # Figure 4: multiseed comparison
    ms = data["multiseed"]
    model_order = ["null[ocv25_qnom]", "vanilla_lstm", "hard_coulomb_lstm", "hard_coulomb_tcn", "hc_anchor_pooled", "hc_anchor_last"]
    labels_map = ["Null", "Vanilla", "HC first", "HC-TCN", "HC pooled", "anchor_last"]
    a_vals, b_vals, a_err, b_err = [], [], [], []
    for m in model_order:
        if m.startswith("null"):
            a_vals.append(13.0297); b_vals.append(7.5278); a_err.append(0); b_err.append(0)
        else:
            ra = next(r for r in ms if r["model"] == m and r["scenario"] == "scenario_A")
            rb = next(r for r in ms if r["model"] == m and r["scenario"] == "scenario_B")
            a_vals.append(float(ra["rmse_pct_mean"])); b_vals.append(float(rb["rmse_pct_mean"]))
            a_err.append(float(ra["rmse_pct_std"])); b_err.append(float(rb["rmse_pct_std"]))
    x = np.arange(len(model_order))
    fig, ax = plt.subplots(figsize=(6.3, 3.1))
    ax.bar(x - 0.18, a_vals, 0.36, yerr=a_err, label="Scenario A", color="#7FB3D5", capsize=2)
    ax.bar(x + 0.18, b_vals, 0.36, yerr=b_err, label="Scenario B", color=colors["ours"], capsize=2)
    ax.set_xticks(x); ax.set_xticklabels(labels_map, rotation=25, ha="right")
    ax.set_ylabel("RMSE (%SOC)")
    ax.legend(frameon=False, ncols=2)
    ax.spines[["top", "right"]].set_visible(False); ax.grid(axis="y", alpha=0.18)
    save_fig(FIG_DIR / "fig4_multiseed.png")

    # Figure 5: recursive policy comparison
    rec = data["recursive"]
    keep = ["windowed_independent", "carried_anchor", "load_gated"]
    vals = [float(next(r for r in rec if r["policy"] == k)["rmse_pct"]) for k in keep]
    n20 = [float(next(r for r in rec if r["policy"] == k)["rmse_n20degC"]) for k in keep]
    fig, ax = plt.subplots(figsize=(4.2, 2.8))
    x = np.arange(len(keep))
    ax.bar(x - 0.17, vals, 0.34, label="Overall", color="#7FB3D5")
    ax.bar(x + 0.17, n20, 0.34, label="−20 °C", color=colors["ours"])
    ax.set_xticks(x); ax.set_xticklabels(["Windowed", "Carried", "Load-gated"], rotation=15, ha="right")
    ax.set_ylabel("RMSE (%SOC)")
    ax.legend(frameon=False)
    ax.spines[["top", "right"]].set_visible(False); ax.grid(axis="y", alpha=0.18)
    save_fig(FIG_DIR / "fig5_recursive.png")

    # Figure 6: eta calibration
    eta = [r for r in data["eta_json"]["rows"] if r["mode"] == "inference_sweep" and r["gamma_mode"] == "nominal"]
    xs = [float(r["eta"]) for r in eta]
    ratios = [float(r["gated_delta_ratio"]) for r in eta]
    rmses = [float(r["rec_rmse_pct"]) for r in eta]
    fig, ax1 = plt.subplots(figsize=(4.2, 2.8))
    ax1.plot(xs, ratios, marker="o", color=colors["hc"], label="Delta ratio")
    ax1.axhline(1.0, color="#555", lw=1, ls="--")
    ax1.set_xlabel("eta")
    ax1.set_ylabel("Delta ratio")
    ax2 = ax1.twinx()
    ax2.plot(xs, rmses, marker="s", color=colors["ours"], label="RMSE")
    ax2.set_ylabel("RMSE (%SOC)")
    ax1.spines["top"].set_visible(False); ax2.spines["top"].set_visible(False)
    ax1.grid(axis="y", alpha=0.18)
    ax1.text(2.02, 1.03, "eta*=2.0", fontsize=8)
    save_fig(FIG_DIR / "fig6_eta.png")

    # Figure 7: EKF vs final
    fig, ax = plt.subplots(figsize=(4.4, 2.8))
    labels = ["HC calibrated", "Best EKF", "EKF low R"]
    vals = [4.4318, 6.8485, 39.1161]
    pvr = [0.0, 5.4755, 33.6695]
    x = np.arange(len(labels))
    ax.bar(x - 0.17, vals, 0.34, label="RMSE", color=[colors["ours"], colors["ekf"], colors["ekf"]])
    ax.bar(x + 0.17, pvr, 0.34, label="PVR", color=[colors["hc"], colors["warn"], colors["warn"]])
    ax.set_xticks(x); ax.set_xticklabels(labels, rotation=15, ha="right")
    ax.set_ylabel("%")
    ax.legend(frameon=False)
    ax.spines[["top", "right"]].set_visible(False); ax.grid(axis="y", alpha=0.18)
    save_fig(FIG_DIR / "fig7_ekf.png")


def refs_from_markdown() -> list[str]:
    md = (ROOT / "JNTETI_Hard_Coulomb_LSTM_Draft_ID.md").read_text(encoding="utf-8")
    return [line.strip() for line in md.splitlines() if re.match(r"^\[\d+\]", line.strip())]


def build():
    data = {
        "dataset": read_csv("results/v5/dataset_variant_comparison.csv"),
        "final": read_csv("results/v5/final_v5_model_comparison.csv"),
        "multiseed": read_csv("results/v5/multiseed/multiseed_summary.csv"),
        "recursive": read_csv("results/v5/recursive_inference/recursive_policy_comparison.csv"),
        "eta": read_csv("results/v5/delta_calibration/eta_gamma_sweep.csv"),
        "eta_json": read_json("results/v5/delta_calibration/eta_gamma_sweep.json"),
        "ekf": read_csv("results/v5/ekf_ecm/recursive_vs_ekf_comparison.csv"),
        "claims": read_json("reports/v5_campaign/claims_register_v2.json"),
        "gate": read_json("reports/v5_campaign/phase10_readiness_gate.json"),
    }
    make_figures(data)

    final = data["final"]
    ms = data["multiseed"]
    rec = data["recursive"]
    eta = data["eta"]
    ekf = data["ekf"]

    def frow(model, scenario=None, policy=None):
        for r in final:
            if r["model"] == model and (scenario is None or r["scenario"] == scenario) and (policy is None or r["inference_policy"] == policy):
                return r
        return {}

    def mrow(model, scenario):
        return next(r for r in ms if r["model"] == model and r["scenario"] == scenario)

    anchor_a = mrow("hc_anchor_last", "scenario_A")
    anchor_b = mrow("hc_anchor_last", "scenario_B")
    original_b = mrow("hard_coulomb_lstm", "scenario_B")
    load_gated = next(r for r in rec if r["policy"] == "load_gated")
    eta2 = next(r for r in eta if r["mode"] == "inference_sweep" and r["eta"] == "2.0" and r["gamma_mode"] == "nominal")
    eta2_rec = next(r for r in data["eta_json"]["rows"] if r["mode"] == "inference_sweep" and r["eta"] == 2.0 and r["gamma_mode"] == "nominal")
    best_ekf = next(r for r in ekf if r["model"] == "EKF_1RC_cont[R=0.01]")

    doc = Document(TEMPLATE)
    clear_doc(doc)
    sec = doc.sections[0]
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.3); sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(1.3); sec.right_margin = Cm(1.3)
    set_columns(sec, None)

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Estimasi SOC Anchor-Aware Hard-Coulomb untuk Baterai Li-Ion")
    run_font(r, "Helvetica", 20, True)
    p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Abyan Hisyam Al'ammar"); run_font(r, "Helvetica", 9)
    add_par(doc, "Program Studi Informatika, Universitas AMIKOM Yogyakarta, Yogyakarta, Indonesia", size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_par(doc, "Received: DD MM YY, Revised: DD MM YY, Accepted: DD MM YY", size=8, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_par(doc, "Corresponding Author: Abyan Hisyam Al'ammar (email: [ISI EMAIL])", size=8, align=WD_ALIGN_PARAGRAPH.LEFT)

    abstract = (
        "ABSTRAK - Estimasi state of charge (SOC) merupakan fungsi utama battery management system (BMS) karena nilai SOC "
        "menentukan batas operasi, proteksi charge-discharge, dan interpretasi kondisi baterai pada kendaraan listrik. "
        "Pada suhu ekstrem, hubungan tegangan, arus, temperatur, dan SOC menjadi lebih ambigu sehingga estimator berbasis "
        "jendela dapat mengalami cold-start anchor berulang. Selain itu, model neural sekuens yang akurat secara RMSE belum "
        "tentu menghasilkan lintasan SOC yang konsisten terhadap arah arus terukur. Studi ini mengevaluasi estimator "
        "Hard-Coulomb anchor-aware dengan inferensi carried terkalibrasi melalui pipeline v5 yang memperbaiki artifact label "
        "dan decimation, ablation bertahap, evaluasi multi-seed, serta pembandingan terhadap OCV+Coulomb dan EKF/1RC kontinu. "
        "Sistem akhir, anchor_last + calibrated carried inference, dipilih karena setiap tahap ablation menutup mode gagal "
        "yang terukur: clamp gagal, anchor-first terbatas, dan delta path perlu kalibrasi laju. "
        "Hasil menunjukkan bahwa desain anchor_last mencapai RMSE windowed Scenario A "
        f"{metric(anchor_a['rmse_pct_mean'], anchor_a['rmse_pct_std'], 2)}% dan Scenario B "
        f"{metric(anchor_b['rmse_pct_mean'], anchor_b['rmse_pct_std'], 2)}%. Pada inference recursive, kalibrasi eta*=2,0 "
        f"memulihkan laju delta dan menurunkan RMSE carried menjadi {fmt(eta2_rec['rec_rmse_pct'],2)}%, termasuk "
        f"{fmt(eta2_rec['rec_rmse_n20'],2)}% pada −20 °C. Baseline EKF terbaik yang diuji mencapai "
        f"{fmt(best_ekf['rmse_pct'],2)}% RMSE. Klaim utama dibatasi pada konsistensi tanda terhadap arus terukur dan protokol "
        "evaluasi v5; sensor fault, kepatuhan keselamatan fungsional, dan superioritas universal tidak diklaim."
    )
    add_par(doc, abstract, style="IEEE Abtract", size=8)
    add_par(doc, "KATA KUNCI - Estimasi SOC, Baterai Li-Ion, Hard-Coulomb, LSTM, Anchor-Aware, Recursive Inference, EKF, Suhu Ekstrem.", style="IEEE Abtract", size=8)

    doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(doc.sections[-1], 2)

    # Introduction
    add_heading(doc, "I. PENDAHULUAN", 1)
    intro = [
        "Estimasi state of charge (SOC) adalah komponen inti BMS karena SOC tidak dapat diukur langsung, tetapi memengaruhi estimasi jarak tempuh, pembatasan daya, strategi pengisian, dan perlindungan baterai. Pada kendaraan listrik, kesalahan SOC dapat menyebabkan keputusan kontrol yang terlalu agresif atau terlalu konservatif, sehingga estimator harus dievaluasi tidak hanya dari galat rata-rata, tetapi juga dari perilaku lintasan.",
        "Suhu ekstrem memperumit estimasi SOC karena resistansi internal, polarisasi, kapasitas efektif, dan respons tegangan berubah terhadap temperatur. Pada suhu rendah, tegangan terminal dapat turun akibat beban dan polarisasi, sehingga estimator yang mengandalkan tegangan awal jendela dapat menafsirkan sel sebagai lebih kosong daripada keadaan sebenarnya.",
        "Model sekuens neural seperti LSTM, TCN, dan CNN-LSTM berguna karena mampu mempelajari hubungan nonlinear antara tegangan, arus, temperatur, dan SOC [6], [7], [13]-[16], [20]. Akan tetapi, model tersebut tidak secara alami menegakkan konsistensi lintasan. Prediksi yang memiliki RMSE moderat masih dapat menaikkan SOC pada saat discharge atau menurunkan SOC pada saat charge.",
        "Masalah lain muncul pada protokol windowed. Setiap jendela memerlukan anchor SOC awal; ketika jendela dimulai saat kondisi cold-load, anchor dapat salah dan kesalahan itu berulang pada banyak jendela. Karena itu, inference deployment perlu dipandang sebagai masalah recursive state carry, bukan hanya regresi jendela independen.",
        "EKF dan ECM tetap menjadi pembanding penting karena keduanya memodelkan estimasi SOC sebagai proses rekursif dengan koreksi tegangan. Namun, koreksi tersebut bergantung pada parameter OCV/ECM dan measurement-noise trust. Pada kondisi suhu ekstrem, koreksi tegangan dapat memperbaiki drift, tetapi juga dapat mendorong lintasan melawan arah arus terukur.",
        "Celah penelitian yang dituju adalah evaluasi bersama antara akurasi, konsistensi lintasan, observabilitas anchor, perilaku recursive inference, dan kalibrasi laju delta. Banyak studi berfokus pada RMSE, sedangkan penelitian ini menanyakan mengapa suatu estimator gagal dan intervensi mana yang memperbaiki mode gagal tersebut.",
    ]
    for t in intro:
        add_par(doc, t)
    add_par(doc, "Kontribusi penelitian ini adalah sebagai berikut.")
    contributions = [
        "Studi ini mengusulkan estimator SOC Hard-Coulomb anchor-aware yang secara struktural menegakkan konsistensi tanda terhadap arus terukur.",
        "Studi ini memberikan bukti ablation bertahap yang menunjukkan mengapa post-hoc clamp, baseline neural vanilla, dan Hard-Coulomb anchor-first belum memadai.",
        "Studi ini memperkenalkan pipeline evaluasi v5 terkoreksi yang menangani artifact label dan decimation pada representasi data yang digunakan.",
        "Studi ini menunjukkan bahwa calibrated carried inference memulihkan fidelity laju delta dan memperbaiki estimasi SOC recursive.",
        "Studi ini membandingkan metode akhir terhadap null OCV+Coulomb dan baseline EKF/1RC kontinu, sehingga frontier correction-vs-consistency dapat diamati.",
    ]
    for i, t in enumerate(contributions, 1):
        add_par(doc, f"{i}. {t}", align=WD_ALIGN_PARAGRAPH.LEFT)

    # Method
    add_heading(doc, "II. METODE", 1)
    add_heading(doc, "A. DATASET DAN KOREKSI V5", 2)
    add_par(doc, "Dataset legacy v4 menyimpan dua sumber bias yang memengaruhi interpretasi hasil: label awal segmen dapat terkontaminasi loaded-start ohmic bias, dan first-sample decimation dapat membuat arus representatif per detik tidak selaras dengan dinamika intradetik. Kampanye v5 membangun empat varian untuk mengisolasi koreksi label dan decimation, lalu memilih v5c sebagai dataset final.")
    add_par(doc, "Varian v5c memakai label ohmic-corrected dan mean-per-second decimation. Koreksi ini mengurangi konflik routing pada representasi per-detik yang digunakan dan mempertahankan jumlah segmen serta split jendela agar perbandingan tetap auditable. Koreksi ohmic adalah batas bawah; komponen polarisasi tidak dimodelkan penuh dan tetap menjadi keterbatasan label.")
    ds_rows = []
    for r in data["dataset"]:
        if r["scenario"] == "A":
            ds_rows.append([r["variant"], r["label_mode"], r["decimation_mode"], r["segments"], f"{r['train_windows']}/{r['val_windows']}/{r['test_windows']}", fmt(r["label_shift_vs_v4_mean_pct"], 2)])
    add_table(doc, "Tabel I. Varian Dataset v4-v5 pada Scenario A.", ["Var.", "Label", "Decimation", "Seg.", "Tr/Val/Test", "Shift"], ds_rows)

    add_heading(doc, "B. FORMULASI SOC DAN COULOMB COUNTING", 2)
    add_par(doc, "Konvensi arus mengikuti pipeline: discharge bernilai negatif dan charge bernilai positif. SOC dinyatakan dalam rentang 0 sampai 1. Untuk arus terukur I_k, interval sampling Δt, kapasitas efektif Q_eff(T), dan faktor laju eta, pembaruan Coulomb counting ditulis sebagai berikut.")
    add_eq(doc, "SOC_t = SOC_0 − (η Δt / (3600 Q_eff(T))) Σ_{k=1..t} I_k", 1)
    add_par(doc, "SOC_t adalah SOC pada timestep t, SOC_0 adalah anchor awal, dan Q_eff(T) menyatakan kapasitas efektif pada temperatur T. Persamaan ini menjadi dasar interpretasi laju perubahan SOC.")

    add_heading(doc, "C. HARD-COULOMB CONSTRAINT", 2)
    add_par(doc, "Hard-Coulomb membatasi tanda delta SOC dengan aturan arus terukur. Mask discharge dan charge didefinisikan oleh ambang arus I_ε.")
    add_eq(doc, "m_t^dis = 1[I_t < −I_ε],    m_t^chg = 1[I_t > I_ε]", 2)
    add_eq(doc, "Δŝ_t = −|Δŝ_t| jika I_t < −I_ε;  +|Δŝ_t| jika I_t > I_ε;  0 atau drift terbatas jika |I_t| ≤ I_ε", 3)
    add_eq(doc, "ŝ_t = ŝ_0 + Σ_{k=1..t} Δŝ_k", 4)
    add_par(doc, "Dengan konstruksi ini, PVR 0,00% adalah sifat struktural terhadap aturan arus terukur, bukan prestasi empiris model. Preconditions-nya penting: arus terukur diasumsikan reliabel, sensor fault tidak dicakup, magnitude delta tidak sepenuhnya dijamin tepat, dan kalibrasi eta tetap diperlukan untuk rate fidelity.")

    add_heading(doc, "D. ANCHOR-AWARE MODEL", 2)
    add_par(doc, "Hard-Coulomb awal memakai anchor-first, yaitu anchor SOC diprediksi dari hidden state awal. Pada kondisi cold-load, hidden state awal dapat didominasi tegangan terpolarisasi sehingga anchor tidak observabel dengan baik. Anchor-last membaca representasi h_T setelah encoder melihat seluruh jendela kausal, lalu anchor diremap ke interval feasible dari akumulasi delta.")
    add_eq(doc, "ŝ_0 = f_θ(h_T) = f_θ(Encoder(x_1:T))", 5)
    add_par(doc, "Desain ini tidak menghilangkan seluruh ambiguitas suhu dingin, tetapi memberi anchor informasi window-level yang tidak tersedia pada anchor-first.")

    add_heading(doc, "E. CALIBRATED CARRIED INFERENCE", 2)
    add_par(doc, "Pada independent windowing, setiap jendela memulai estimasi dari anchor baru. Pada carried inference, SOC hasil jendela sebelumnya menjadi state awal jendela berikutnya sehingga artifact cold-start tidak berulang. Kebijakan load-gated melakukan re-anchor hanya ketika kondisi rest atau low-load membuat anchor tegangan lebih dapat dipercaya.")
    add_eq(doc, "ŝ_t^rec = ŝ_{t−1}^rec + η* Δŝ_t", 6)
    add_eq(doc, "r_Δ = mean(|Δŝ|) / mean(|Δs|)", 7)
    add_par(doc, f"Pada bukti v5 yang tersedia, eta*=2,0 dipilih karena memulihkan r_Δ dari 0,751 menjadi {fmt(eta2['gated_delta_ratio'],3)}. Dengan carried inference, konfigurasi ini menghasilkan RMSE {fmt(eta2_rec['rec_rmse_pct'],2)}%. Bukti ini berasal dari Scenario A seed 42 dan harus diperlakukan sebagai single-checkpoint/single-scenario sampai re-derivation validation-chain dan konfirmasi multi-seed selesai. Inilah konfigurasi akhir anchor_last + calibrated carried inference.")

    add_heading(doc, "F. BASELINE METHODS", 2)
    add_par(doc, "Baseline yang diuji meliputi null OCV+Coulomb, Vanilla LSTM/TCN, post-hoc clamp, Hard-Coulomb anchor-first, HC-TCN, anchor_last, anchor_pooled, dan EKF kontinu berbasis OCV-Rint serta 1RC-ECM. EKF memakai parameter literature-like, bukan identifikasi khusus sel ini; karena itu EKF dipakai sebagai pembanding klasik yang transparan, bukan batas performa mutlak.")
    add_figure(doc, FIG_DIR / "fig1_architecture.png", "Gambar 1. Sistem akhir memakai anchor_last untuk observabilitas jendela dan carried inference terkalibrasi untuk menjaga kontinuitas recursive.", width=3.15)

    # Experiment setup
    add_heading(doc, "III. SETUP EKSPERIMEN", 1)
    add_par(doc, "Scenario A mengevaluasi temperature-OOD, sedangkan Scenario B memakai chronological in-distribution split. Model windowed utama dievaluasi dengan seeds 1-5. Eksperimen recursive policy, eta calibration, dan EKF kontinu memakai Scenario A seed 42 karena bertujuan memetakan perilaku deployment-like tanpa rerun training berat.")
    add_eq(doc, "RMSE = sqrt((1/N) Σ_{i=1..N} (s_i − ŝ_i)^2)", 8)
    add_eq(doc, "MAE = (1/N) Σ_{i=1..N} |s_i − ŝ_i|", 9)
    add_eq(doc, "MaxE = max_i |s_i − ŝ_i|", 10)
    add_eq(doc, "PVR_dis = Σ_t 1[I_t < −I_ε] 1[Δŝ_t > ε] / Σ_t 1[I_t < −I_ε]", 11)
    add_par(doc, "RMSE dan MAE mengukur galat rata-rata, MaxE mengukur galat terburuk, PVR_dis mengukur pelanggaran arah saat discharge, dan r_Δ mengukur fidelity magnitude delta.")

    # Results
    add_heading(doc, "IV. HASIL DAN PEMBAHASAN", 1)
    add_heading(doc, "A. BAGAIMANA KEGAGALAN AWAL MENGARAH KE METODE AKHIR?", 2)
    add_figure(doc, FIG_DIR / "fig2_evolution.png", "Gambar 2. Alur failure analysis menunjukkan bahwa sistem akhir muncul dari rangkaian koreksi mode gagal, bukan pemilihan model arbitrer.", width=3.15)
    add_par(doc, "Rasional pemilihan sistem akhir adalah rangkaian eliminasi mode gagal. Windowed inference memberi cold-start anchor berulang; vanilla model tidak menjamin lintasan; penalty-based method tidak mengubah ruang keluaran; post-hoc clamp hanya menyaring keluaran setelah prediksi; Hard-Coulomb awal memperbaiki tanda tetapi tetap anchor-limited; v5c memperbaiki artifact label; anchor_last memperbaiki observabilitas; carried inference mengurangi cold-start; dan eta calibration memulihkan laju delta.")
    failure_rows = [
        ["Windowed cold-start", "Cold −20 °C 16.75", "carried inference", "kontinuitas recursive"],
        ["Physics-blind vanilla", "lintasan tidak terikat", "Hard-Coulomb", "sign consistency"],
        ["Post-hoc clamp", "25.03/30.07 RMSE", "training-through-constraint", "representasi terikat"],
        ["Anchor-first", "B gagal 5/5 seeds", "anchor_last", "observabilitas anchor"],
        ["Delta under-rate", "r_Δ=0.751", "eta*=2.0", "r_Δ=1.002"],
        ["EKF inconsistency", "PVR 5-40%", "calibrated HC", "alternatif konsisten"],
    ]
    add_table(doc, "Tabel II. Rasional Pemilihan Sistem Akhir.", ["Failure Mode", "Evidence", "Fix", "Final Role"], failure_rows, size=6.0)

    add_heading(doc, "B. APAKAH KOREKSI DATASET V5 MENGUBAH KESIMPULAN?", 2)
    add_figure(doc, FIG_DIR / "fig3_dataset_shift.png", "Gambar 3. Koreksi v5c mengubah magnitudo label tanpa mengubah kesimpulan utama bahwa anchor-aware HC tetap unggul pada protokol v5.", width=3.05)
    add_par(doc, "Koreksi v5c membuat narasi v4 lebih jujur: sebagian MaxE ekstrem v4 berasal dari artifact label. Namun, klaim inti justru lebih kuat setelah koreksi, karena anchor_last tetap mengalahkan null dan vanilla pada kedua scenario, sementara kelemahan original HC di Scenario B menjadi jelas.")

    add_heading(doc, "C. APAKAH CONSTRAINT PERLU DILATIH END-TO-END?", 2)
    add_par(doc, "Post-hoc clamp pada keluaran vanilla menghasilkan RMSE 25,03% di Scenario A dan 30,07% di Scenario B. Clamp dapat memaksa tanda setelah prediksi, tetapi tidak memperbaiki representasi lintasan yang dipelajari model bebas. Hasil ini mendukung training-through-constraint: constraint harus menjadi bagian dari forward computation yang dipelajari.")

    add_heading(doc, "D. APAKAH ANCHOR-LAST STABIL?", 2)
    main_rows = [
        ["Null OCV+Coulomb", fmt(frow("null[ocv25_qnom]", "scenario_A")["rmse_pct"]), fmt(frow("null[ocv25_qnom]", "scenario_B")["rmse_pct"]), "referensi"],
        ["Vanilla LSTM", metric(mrow("vanilla_lstm", "scenario_A")["rmse_pct_mean"], mrow("vanilla_lstm", "scenario_A")["rmse_pct_std"]), metric(mrow("vanilla_lstm", "scenario_B")["rmse_pct_mean"], mrow("vanilla_lstm", "scenario_B")["rmse_pct_std"]), "tanpa constraint"],
        ["Post-hoc clamp", fmt(frow("vanilla+posthoc_clamp", "scenario_A")["rmse_pct"]), fmt(frow("vanilla+posthoc_clamp", "scenario_B")["rmse_pct"]), "collapse"],
        ["Original HC", metric(mrow("hard_coulomb_lstm", "scenario_A")["rmse_pct_mean"], mrow("hard_coulomb_lstm", "scenario_A")["rmse_pct_std"]), metric(mrow("hard_coulomb_lstm", "scenario_B")["rmse_pct_mean"], mrow("hard_coulomb_lstm", "scenario_B")["rmse_pct_std"]), "anchor-first"],
        ["anchor_last", metric(anchor_a["rmse_pct_mean"], anchor_a["rmse_pct_std"]), metric(anchor_b["rmse_pct_mean"], anchor_b["rmse_pct_std"]), "selected windowed"],
        ["HC carried eta*=2.0", fmt(eta2_rec["rec_rmse_pct"]), "n/a", "recursive"],
        ["Best EKF 1RC", fmt(best_ekf["rmse_pct"]), "n/a", "continuous"],
    ]
    add_table(doc, "Tabel III. Perbandingan Utama Model dan Baseline.", ["Model", "A RMSE", "B RMSE", "Peran"], main_rows, size=5.9)
    multi_rows = []
    for model in ["vanilla_lstm", "hard_coulomb_lstm", "hc_anchor_last", "hc_anchor_pooled", "hard_coulomb_tcn"]:
        a = mrow(model, "scenario_A"); b = mrow(model, "scenario_B")
        multi_rows.append([model.replace("hard_coulomb_lstm", "HC first").replace("hc_anchor_", "anchor_"), metric(a["rmse_pct_mean"], a["rmse_pct_std"]), metric(b["rmse_pct_mean"], b["rmse_pct_std"]), metric(a["maxe_pct_mean"], a["maxe_pct_std"])])
    add_table(doc, "Tabel IV. Hasil Multi-Seed v5c.", ["Model", "A RMSE", "B RMSE", "A MaxE"], multi_rows, size=5.9)
    add_figure(doc, FIG_DIR / "fig4_multiseed.png", "Gambar 4. Anchor_last memberi RMSE rata-rata terbaik pada Scenario A dan B, sedangkan original HC gagal sistematis di Scenario B.", width=3.15)
    add_par(doc, f"Anchor_last mencapai {metric(anchor_a['rmse_pct_mean'], anchor_a['rmse_pct_std'],2)}% pada Scenario A dan {metric(anchor_b['rmse_pct_mean'], anchor_b['rmse_pct_std'],2)}% pada Scenario B. Original HC Scenario B berada pada {metric(original_b['rmse_pct_mean'], original_b['rmse_pct_std'],2)}% dan laporan v5 menyatakan kegagalan sistematis pada 5/5 seeds. Klaim stabilitas dibuat pada level keluarga anchor-aware, bukan kemenangan tunggal setiap seed.")

    add_heading(doc, "E. APAKAH RECURSIVE INFERENCE MENGURANGI COLD-START ARTIFACT?", 2)
    rec_rows = [
        ["Windowed", fmt(next(r for r in rec if r["policy"] == "windowed_independent")["rmse_pct"]), fmt(next(r for r in rec if r["policy"] == "windowed_independent")["rmse_n20degC"]), fmt(next(r for r in rec if r["policy"] == "windowed_independent")["rmse_40degC"]), "cold-start berulang"],
        ["Carried", fmt(next(r for r in rec if r["policy"] == "carried_anchor")["rmse_pct"]), fmt(next(r for r in rec if r["policy"] == "carried_anchor")["rmse_n20degC"]), fmt(next(r for r in rec if r["policy"] == "carried_anchor")["rmse_40degC"]), "cold membaik, warm drift"],
        ["Load-gated", fmt(load_gated["rmse_pct"]), fmt(load_gated["rmse_n20degC"]), fmt(load_gated["rmse_40degC"]), "gate terbaik"],
        ["Eta*=2.0", fmt(eta2_rec["rec_rmse_pct"]), fmt(eta2_rec["rec_rmse_n20"]), fmt(eta2_rec["rec_rmse_40"]), "rate fidelity"],
    ]
    add_table(doc, "Tabel V. Recursive Inference dan Eta Calibration.", ["Policy", "RMSE", "−20 °C", "40 °C", "Interpretasi"], rec_rows, size=5.9)
    add_figure(doc, FIG_DIR / "fig5_recursive.png", "Gambar 5. Recursive policy mengurangi artifact cold-start, tetapi carried inference perlu gate atau kalibrasi agar tidak menambah drift hangat.", width=3.05)
    add_par(doc, f"Load-gated mencapai RMSE {fmt(load_gated['rmse_pct'],2)}% dan −20 °C {fmt(load_gated['rmse_n20degC'],2)}%. Pure carried menurunkan error cold tetapi menaikkan error 40 °C, sehingga recursive inference perlu dikaitkan dengan kondisi fisik dan kalibrasi delta.")

    add_heading(doc, "F. APAKAH ETA CALIBRATION MEMPERBAIKI DELTA-PATH?", 2)
    add_figure(doc, FIG_DIR / "fig6_eta.png", "Gambar 6. Kalibrasi eta memulihkan rasio delta dari 0,751 menjadi 1,002 dan menurunkan RMSE carried recursive hingga 4,43% pada protokol v5.", width=3.05)
    add_par(doc, "Hard-Coulomb awal sudah benar dalam arah, tetapi belum tentu benar dalam laju. Delta ratio 0,751 berarti model mengestimasi gerak SOC terlalu kecil relatif terhadap delta label. Dengan eta*=2,0, rasio menjadi 1,002; recursive carried mencapai 4,43% dan −20 °C 3,59%. Retraining pada eta yang sama tidak self-calibrate karena head magnitude mengompensasi envelope.")

    add_heading(doc, "G. BAGAIMANA DIBANDINGKAN DENGAN EKF?", 2)
    ekf_rows = [
        ["Calibrated HC", "4.43", "3.59", "0.00", "single checkpoint/scenario"],
        ["Best EKF 1RC", fmt(best_ekf["rmse_pct"]), fmt(best_ekf["rmse_n20"]), fmt(best_ekf["pvr_disch_eps0"]), "parameter literature-like"],
        ["EKF low R", "39.12", "36.06", "33.67", "over-trust voltage"],
    ]
    add_table(doc, "Tabel VI. EKF dan Sistem Akhir.", ["Metode", "RMSE", "−20 °C", "PVR", "Limitasi"], ekf_rows, size=5.9)
    add_figure(doc, FIG_DIR / "fig7_ekf.png", "Gambar 7. Dibanding EKF terbaik yang diuji, calibrated HC memiliki RMSE lebih rendah dan mempertahankan PVR 0,00% by construction.", width=3.05)
    add_par(doc, "EKF adalah pembanding rekursif yang wajar, tetapi hasilnya sensitif terhadap R. EKF 1RC terbaik mencapai RMSE 6,85% dan PVR 5,48%; konfigurasi R kecil mencapai RMSE 39,12% dan PVR 33,67%. Karena parameter EKF tidak diidentifikasi khusus untuk sel ini, hasil ini tidak membuktikan superioritas universal, melainkan menunjukkan frontier correction-vs-consistency pada konfigurasi yang diuji.")

    add_heading(doc, "H. APA BATASAN KLAIM?", 2)
    add_par(doc, "Klaim utama adalah measured-current sign consistency, bukan kebenaran fisik penuh. Jika sensor arus salah, aturan Hard-Coulomb mengikuti arus yang salah. Eta*=2,0 masih perlu diturunkan dari validation chains dan dikonfirmasi pada Scenario B serta multi-seed. EKF memakai parameter literature-like; ECM teridentifikasi cell-specific dapat mengubah gap akurasi. Deployment edge tidak diklaim karena belum ada validasi WCET, RAM, dan accumulator kuantisasi.")

    # Conclusion and end matter
    add_heading(doc, "V. KESIMPULAN", 1)
    add_par(doc, "Sistem akhir anchor_last + calibrated carried inference ditemukan melalui staged failure analysis. Post-hoc clamp menunjukkan bahwa constraint harus dilatih melalui forward computation. Anchor_last memperbaiki observabilitas jendela, carried inference mengurangi cold-start berulang, dan eta calibration memperbaiki underestimation delta-path. Dalam protokol v5 yang diuji, anchor_last + calibrated carried inference mencapai 4,43% RMSE dan −20 °C 3,59%, lebih rendah daripada EKF terbaik yang diuji sebesar 6,85%. Pekerjaan berikutnya mencakup validasi eta pada validation chains, konfirmasi Scenario B dan multi-seed, diagnostik sensor fault, accumulator sign-preserving untuk kuantisasi, serta evaluasi pada dataset baterai yang lebih luas.")

    for title, text in [
        ("KONFLIK KEPENTINGAN", "Penulis menyatakan tidak terdapat konflik kepentingan."),
        ("KONTRIBUSI PENULIS", "Abyan Hisyam Al'ammar: konseptualisasi, metodologi, perangkat lunak, validasi, analisis formal, investigasi, kurasi data, visualisasi, penulisan draf awal, dan penyuntingan naskah. [ISI NAMA DOSEN PEMBIMBING]: supervisi, validasi metodologi, dan penelaahan naskah."),
        ("UCAPAN TERIMA KASIH", "Penulis mengucapkan terima kasih kepada [ISI NAMA DOSEN/PROGRAM STUDI/LAB] atas arahan dan masukan selama proses penelitian."),
    ]:
        add_heading(doc, title, 1)
        add_par(doc, text)

    add_heading(doc, "REFERENSI", 1)
    for ref in refs_from_markdown()[:24]:
        p = doc.add_paragraph(style="IEEE Reference Item")
        r = p.add_run(ref)
        run_font(r, size=7.4)
        p.paragraph_format.space_after = Pt(2)

    add_heading(doc, "LAMPIRAN", 1)
    add_par(doc, "Lampiran ini merangkum bukti reproduksibilitas yang tidak dimuat penuh pada teks utama. Final ablation matrix berisi 62 baris; claims register v2 berisi 17 klaim; readiness gate Phase 10 melaporkan 52/52 PASS. Notebook pendukung berada pada notebooks/ablation_studies_v5_final/12-19.")
    phase_rows = [["P0", "legacy freeze"], ["P1", "dataset v5c"], ["P2", "headline retrain"], ["P3", "multi-seed"], ["P4", "recursive policy"], ["P5", "eta calibration"], ["P6", "EKF continuous"], ["P7", "final matrix"], ["P8", "claims register"], ["P9", "final report"], ["P10", "52/52 gate"]]
    add_table(doc, "Tabel VII. Ringkasan Fase Kampanye v5.", ["Fase", "Isi"], phase_rows, size=6.2)
    claim_rows = [[str(c["id"]), c["status"], c["claim"][:70]] for c in data["claims"]["claims"] if c["id"] in (1, 2, 5, 6, 11, 12, 15, 16, 17)]
    add_table(doc, "Tabel VIII. Ringkasan Klaim dan Batasan.", ["ID", "Status", "Klaim"], claim_rows, size=6.0)

    doc.save(OUT)
    Document(OUT)
    with zipfile.ZipFile(OUT) as z:
        assert "word/document.xml" in set(z.namelist())

    text = "\n".join(p.text for p in Document(OUT).paragraphs)
    assert "[REF NEEDED]" not in text
    assert "[PERLU DIVERIFIKASI]" not in text
    for bad in [r"\frac", r"\sum", r"\mathrm", r"\hat", r"\mathbb", "+/-", "-20 C", "`"]:
        assert bad not in text, bad

    trace = [
        "# Traceability v2 - JNTETI SOC Hard-Coulomb",
        "",
        "## Artifact Paths Used",
        "- reports/v5_campaign/phase1_dataset_v5_report.md",
        "- reports/v5_campaign/phase3_multiseed_report.md",
        "- reports/v5_campaign/phase4_recursive_policies_report.md",
        "- reports/v5_campaign/phase5_delta_calibration_report.md",
        "- reports/v5_campaign/phase6_ekf_baselines_report.md",
        "- reports/v5_campaign/phase7_final_ablation_model_card.md",
        "- reports/v5_campaign/phase9_final_v5_report.md",
        "- reports/v5_campaign/phase10_manuscript_readiness_gate.md",
        "- reports/v5_campaign/claims_register_v2.md/json",
        "- results/v5/dataset_variant_comparison.csv",
        "- results/v5/final_v5_model_comparison.csv",
        "- results/v5/final_ablation_matrix.csv",
        "- results/v5/multiseed/multiseed_summary.csv",
        "- results/v5/recursive_inference/recursive_policy_comparison.csv",
        "- results/v5/delta_calibration/eta_gamma_sweep.csv/json",
        "- results/v5/ekf_ecm/recursive_vs_ekf_comparison.csv",
        "- notebooks/ablation_studies_v5_final/README.md",
        "",
        "## Figures Used",
        "- drafts/figures_v2/fig1_architecture.png",
        "- drafts/figures_v2/fig2_evolution.png",
        "- drafts/figures_v2/fig3_dataset_shift.png",
        "- drafts/figures_v2/fig4_multiseed.png",
        "- drafts/figures_v2/fig5_recursive.png",
        "- drafts/figures_v2/fig6_eta.png",
        "- drafts/figures_v2/fig7_ekf.png",
        "",
        "## Tables Used",
        "- Table I dataset variants from dataset_variant_comparison.csv",
        "- Table II failure-mode synthesis from phase reports and claims register",
        "- Table III main comparison from final_v5_model_comparison.csv, multiseed_summary.csv, eta sweep, and EKF comparison",
        "- Table IV multi-seed from multiseed_summary.csv",
        "- Table V recursive/eta from recursive_policy_comparison.csv and eta_gamma_sweep.csv",
        "- Table VI EKF vs final from recursive_vs_ekf_comparison.csv",
        "- Appendix Tables VII-VIII from phase reports and claims_register_v2.json",
        "",
        "## Numbers Used",
        f"- anchor_last Scenario A: {metric(anchor_a['rmse_pct_mean'], anchor_a['rmse_pct_std'], 3)}",
        f"- anchor_last Scenario B: {metric(anchor_b['rmse_pct_mean'], anchor_b['rmse_pct_std'], 3)}",
        f"- original HC Scenario B: {metric(original_b['rmse_pct_mean'], original_b['rmse_pct_std'], 3)}",
        f"- load_gated RMSE: {fmt(load_gated['rmse_pct'], 3)}",
        f"- eta*=2.0 gated RMSE: {fmt(eta2['gated_rmse_pct'], 4)}, −20 °C: {fmt(eta2['gated_rmse_n20'], 4)}, delta ratio: {fmt(eta2['gated_delta_ratio'], 4)}",
        f"- eta*=2.0 carried recursive RMSE: {fmt(eta2_rec['rec_rmse_pct'], 4)} and −20 °C: {fmt(eta2_rec['rec_rmse_n20'], 4)} from eta_gamma_sweep.json",
        f"- best EKF RMSE: {fmt(best_ekf['rmse_pct'], 4)}, PVR: {fmt(best_ekf['pvr_disch_eps0'], 4)}",
        "",
        "## References Changed",
        "- Removed manuscript [REF NEEDED] placeholder.",
        "- Retained 24 references from existing repo draft.",
        "- References were not newly fetched; bibliography requires manual verification before journal submission.",
        "",
        "## Unresolved Placeholders",
        "- [ISI EMAIL]",
        "- [ISI NAMA DOSEN PEMBIMBING]",
        "- [ISI NAMA DOSEN/PROGRAM STUDI/LAB]",
        "",
        "## Missing Artifacts / Assumptions",
        "- `reports/v5_campaign/manuscript_rewrite_brief.md` not present; Phase 9 rewrite brief section used.",
        "- `results/v5/final_figures/` and `results/v5/final_tables/` not present; generated v2 figures from CSV/JSON and used existing CSV/JSON tables.",
        "- eta*=2.0 marked as single-checkpoint/single-scenario pending validation-chain and multi-seed confirmation.",
    ]
    TRACE.write_text("\n".join(trace) + "\n", encoding="utf-8")

    notes = [
        "# Revision Notes v2",
        "",
        "## Rewritten",
        "- Abstract rewritten to 200-250 word target and reduced to core results.",
        "- Introduction rewritten into scholarly argument flow with research contributions.",
        "- Method rewritten for formal sign convention, limitations, and clearer anchor/recursive logic.",
        "- Results reorganized by scientific questions rather than artifact order.",
        "- Added mandatory Rasional Pemilihan Sistem Akhir failure-mode table.",
        "",
        "## Moved to Appendix",
        "- P0-P10 phase checklist.",
        "- Claims register summary.",
        "- 62-row ablation matrix and notebook evidence mention.",
        "",
        "## Formatting Fixes",
        "- Removed raw LaTeX commands from equations.",
        "- Replaced backticks, +/- notation, and -20 C notation.",
        "- Removed [REF NEEDED] and [PERLU DIVERIFIKASI] from manuscript body.",
        "- Simplified main tables and regenerated seven 300 ppi figures.",
        "",
        "## Remaining Before Supervisor Submission",
        "- Fill email, supervisor name, and acknowledgment placeholders.",
        "- Supervisor should confirm title preference and Indonesian terminology.",
        "- Confirm whether eta*=2.0 should be described as validation-derived if new validation artifact becomes available.",
        "",
        "## Remaining Before Journal Submission",
        "- Verify every reference and DOI manually/programmatically.",
        "- Run visual DOCX/PDF render in Microsoft Word or LibreOffice.",
        "- Add final dates and author metadata required by JNTETI.",
        "- Re-check table/figure placement after Word field/layout update.",
    ]
    NOTES.write_text("\n".join(notes) + "\n", encoding="utf-8")

    print(OUT)
    print(TRACE)
    print(NOTES)


if __name__ == "__main__":
    build()
