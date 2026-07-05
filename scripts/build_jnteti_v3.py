from __future__ import annotations

import csv
import json
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "Template-JNTETI-2025-ENG (3).docx"
OUT = ROOT / "drafts" / "JNTETI_SOC_Hard_Coulomb_v5_Draft_ID_v3.docx"
TRACE = ROOT / "drafts" / "JNTETI_SOC_Hard_Coulomb_v5_Draft_ID_v3_traceability.md"
NOTES = ROOT / "drafts" / "JNTETI_SOC_Hard_Coulomb_v5_Draft_ID_v3_revision_notes.md"

FIG_DIR = ROOT / "outputs" / "figures"
TABLE_DIR = ROOT / "outputs" / "tables"
ASSET_DIR = ROOT / "outputs" / "manuscript_assets"

TITLE = "Estimasi SOC Anchor-Aware Hard-Coulomb untuk Baterai Li-Ion"
AUTHOR = "Abyan Hisyam Al'ammar"
AFFILIATION = "Program Studi Informatika, Universitas AMIKOM Yogyakarta, Yogyakarta, Indonesia"


def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_columns(section, num: int) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), "403")


def set_page(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(23)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(13)
    section.right_margin = Mm(13)
    set_columns(section, 1)


def set_run_font(run, name="Times New Roman", size=9, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)


def add_para(doc, text="", style=None, align=None, font="Times New Roman", size=9, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    if text:
        r = p.add_run(text)
        set_run_font(r, font, size, bold, italic)
    return p


def add_heading(doc, text, level=1):
    style = {1: "IEEE Heading 1", 2: "IEEE Heading 2", 3: "IEEE Heading 3"}.get(level)
    if style not in [s.name for s in doc.styles]:
        style = "Normal"
    p = add_para(doc, text, style=style, font="Helvetica", size=9, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_bullets(doc, items):
    for item in items:
        p = add_para(doc)
        p.paragraph_format.left_indent = Mm(4)
        p.paragraph_format.first_line_indent = Mm(-4)
        r = p.add_run("• ")
        set_run_font(r, "Times New Roman", 9)
        r = p.add_run(item)
        set_run_font(r, "Times New Roman", 9)


def add_numbered(doc, items):
    for i, item in enumerate(items, 1):
        p = add_para(doc)
        p.paragraph_format.left_indent = Mm(5)
        p.paragraph_format.first_line_indent = Mm(-5)
        r = p.add_run(f"{i}. ")
        set_run_font(r, "Times New Roman", 9)
        r = p.add_run(item)
        set_run_font(r, "Times New Roman", 9)


def add_equation(doc, eq: str, num: int):
    p = add_para(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{eq}    ({num})")
    set_run_font(r, "Cambria Math", 9.5)
    return p


def set_cell_margins(cell, top=70, start=70, bottom=70, end=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_in):
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def fit_table(table, widths, font_size=7.0):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for r_i, row in enumerate(table.rows):
        for c_i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(c_i, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_i == 0:
                shade_cell(cell, "D9EAF7")
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                if c_i == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, "Times New Roman", font_size, bold=(r_i == 0))


def add_table_caption(doc, caption):
    p = add_para(doc, caption, font="Times New Roman", size=8, bold=False)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    return p


def add_dataframe_table(doc, caption, df: pd.DataFrame, widths, font_size=7.0):
    add_table_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for c, col in enumerate(df.columns):
        table.rows[0].cells[c].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for c, col in enumerate(df.columns):
            val = row[col]
            cells[c].text = "" if pd.isna(val) else str(val)
    fit_table(table, widths, font_size=font_size)
    add_para(doc, "")
    return table


def add_figure(doc, path, caption, width_in, full_width=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_in))
    cap = add_para(doc, caption, font="Times New Roman", size=8)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if full_width:
        cap.paragraph_format.space_after = Pt(6)
    return p


def section(doc, columns: int):
    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    set_page(sec)
    set_columns(sec, columns)
    return sec


def roman(n):
    vals = [(10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I")]
    out = ""
    for v, s in vals:
        while n >= v:
            out += s
            n -= v
    return out


def fmt(x, nd=2):
    try:
        return f"{float(x):.{nd}f}"
    except Exception:
        return str(x)


def metric(mean, std=None, nd=2):
    return f"{fmt(mean, nd)} ± {fmt(std, nd)}" if std is not None else fmt(mean, nd)


def load_tables():
    return {p.stem: pd.read_csv(p) for p in TABLE_DIR.glob("table0*.csv")}


def table_dataset(df):
    keep = [
        "variant",
        "scenario",
        "label_mode",
        "decimation_mode",
        "train_windows",
        "val_windows",
        "test_windows",
        "label_shift_vs_v4_mean_pct",
        "leakage_overlaps",
    ]
    out = df[keep].copy()
    out.columns = ["Varian", "Sken.", "Label", "Decim.", "Train", "Val", "Test", "Shift mean", "Leak"]
    out["Shift mean"] = out["Shift mean"].map(lambda x: fmt(x, 2))
    return out


def table_failure(df):
    diag = {
        "Windowed cold-start artifact": "Anchor diulang pada setiap jendela.",
        "Physics-blind vanilla trajectory": "RMSE tidak menjamin arah lintasan.",
        "Post-hoc clamp collapse": "Filter keluaran tidak membentuk representasi.",
        "Anchor-first bottleneck": "Hidden state awal miskin observabilitas.",
        "Delta-path underestimation": "Arah benar, magnitude laju terlalu kecil.",
        "EKF correction inconsistency": "Koreksi tegangan dapat melawan arus.",
    }
    out = df.copy()
    out.insert(2, "diagnosis", out["failure_mode"].map(diag).fillna("Mode gagal terukur."))
    out.columns = ["Mode gagal", "Bukti", "Diagnosis", "Perbaikan", "Peran akhir"]
    return out


def table_model(df):
    out = df.copy()
    out.columns = ["Metode", "RMSE A", "RMSE B", "Catatan"]
    return out


def table_multiseed(df):
    out = df.copy()
    out["RMSE"] = [metric(a, b) for a, b in zip(out.rmse_pct_mean, out.rmse_pct_std)]
    out["MaxE"] = [metric(a, b) for a, b in zip(out.maxe_pct_mean, out.maxe_pct_std)]
    out["−20 °C"] = out.rmse_n20degC_mean.map(lambda x: fmt(x, 2))
    out["40 °C"] = out.rmse_40degC_mean.map(lambda x: fmt(x, 2))
    out = out[["model", "scenario", "RMSE", "MaxE", "−20 °C", "40 °C"]]
    out.columns = ["Model", "Skenario", "RMSE", "MaxE", "RMSE −20 °C", "RMSE 40 °C"]
    return out


def table_recursive(df):
    out = df.copy()
    keep = ["policy", "rmse_pct", "maxe_pct", "delta_ratio_disch", "rmse_n20degC", "rmse_40degC", "reanchor_pct", "source"]
    out = out[keep].copy()
    for c in ["rmse_pct", "maxe_pct", "delta_ratio_disch", "rmse_n20degC", "rmse_40degC", "reanchor_pct"]:
        out[c] = out[c].map(lambda x: fmt(x, 2))
    out.columns = ["Policy", "RMSE", "MaxE", "rΔ", "−20 °C", "40 °C", "Re-anchor", "Sumber"]
    return out


def table_ekf(df):
    out = df.copy()
    for c in ["rmse_pct", "rmse_n20", "pvr_disch_eps0"]:
        out[c] = out[c].map(lambda x: fmt(x, 2))
    out.columns = ["Metode", "RMSE", "−20 °C", "PVR", "Keterbatasan"]
    return out


def table_claims(df):
    out = df.copy()
    out = out[["id", "status", "claim", "evidence"]]
    out.columns = ["ID", "Status", "Klaim", "Bukti"]
    return out


REFERENCES = [
    '[1] N. Nitta, F. Wu, J. T. Lee, and G. Yushin, "Li-ion battery materials: present and future," Materials Today, vol. 18, no. 5, pp. 252-264, May 2015, doi: 10.1016/j.mattod.2014.10.040.',
    '[2] M. A. Hannan, M. S. H. Lipu, A. Hussain, and A. Mohamed, "A review of lithium-ion battery state of charge estimation and management system in electric vehicle applications: Challenges and recommendations," Renewable and Sustainable Energy Reviews, vol. 78, pp. 834-854, Oct. 2017, doi: 10.1016/j.rser.2017.05.001.',
    '[3] K. W. E. Cheng, B. P. Divakar, H. Wu, K. Ding, and H. F. Ho, "Battery-management system (BMS) and SOC development for electrical vehicles," IEEE Transactions on Vehicular Technology, vol. 60, no. 1, pp. 76-88, Jan. 2011, doi: 10.1109/TVT.2010.2089647.',
    '[4] M. Berecibar, I. Gandiaga, I. Villarreal, N. Omar, J. Van Mierlo, and P. Van den Bossche, "Critical review of state of health estimation methods of Li-ion batteries for real applications," Renewable and Sustainable Energy Reviews, vol. 56, pp. 572-587, Apr. 2016, doi: 10.1016/j.rser.2015.11.042.',
    '[5] R. Xiong, J. Cao, Q. Yu, H. He, and F. Sun, "Critical review on the battery state of charge estimation methods for electric vehicles," IEEE Access, vol. 6, pp. 1832-1843, 2018, doi: 10.1109/ACCESS.2017.2780258.',
    '[6] E. Chemali, P. J. Kollmeyer, M. Preindl, and A. Emadi, "State-of-charge estimation of lithium-ion batteries using deep neural networks: A machine learning approach," Journal of Power Sources, vol. 400, pp. 242-255, Oct. 2018, doi: 10.1016/j.jpowsour.2018.06.104.',
    '[7] E. Chemali, P. J. Kollmeyer, M. Preindl, R. Ahmed, and A. Emadi, "Long short-term memory networks for accurate state-of-charge estimation of lithium-ion batteries," IEEE Transactions on Industrial Electronics, vol. 65, no. 8, pp. 6730-6739, Aug. 2018, doi: 10.1109/TIE.2017.2787586.',
    '[8] T. Waldmann, B. I. Hogg, and M. Wohlfahrt-Mehrens, "Li plating as unwanted side reaction in commercial Li-ion cells: A review," Journal of Power Sources, vol. 384, pp. 107-124, Apr. 2018, doi: 10.1016/j.jpowsour.2018.02.063.',
    '[9] M. Petzl and M. A. Danzer, "Nondestructive detection, characterization, and quantification of lithium plating in commercial lithium-ion batteries," Journal of Power Sources, vol. 254, pp. 80-87, May 2014, doi: 10.1016/j.jpowsour.2013.12.060.',
    '[10] S. Ma, N. Jiang, P. Gao, C. Li, and M. Wang, "Temperature effect and thermal impact in lithium-ion batteries: A review," Progress in Natural Science: Materials International, vol. 28, no. 6, pp. 653-666, Dec. 2018, doi: 10.1016/j.pnsc.2018.11.002.',
    '[11] D. N. T. How, M. A. Hannan, M. S. H. Lipu, P. J. Ker, and A. Hussain, "State of charge estimation for lithium-ion batteries using machine learning techniques: A review," IEEE Access, vol. 7, pp. 136116-136136, 2019, doi: 10.1109/ACCESS.2019.2942213.',
    '[12] C. Vidal, P. Malysz, P. J. Kollmeyer, and A. Emadi, "Machine learning applied to electrified vehicle battery state of charge and state of health estimation: State-of-the-art," IEEE Access, vol. 8, pp. 52796-52814, 2020, doi: 10.1109/ACCESS.2020.2980984.',
    '[13] S. Hochreiter and J. Schmidhuber, "Long short-term memory," Neural Computation, vol. 9, no. 8, pp. 1735-1780, Nov. 1997, doi: 10.1162/neco.1997.9.8.1735.',
    '[14] S. Bai, J. Z. Kolter, and V. Koltun, "An empirical evaluation of generic convolutional and recurrent networks for sequence modeling," arXiv:1803.01271, 2018.',
    '[15] C. Lea, M. D. Flynn, R. Vidal, A. Reiter, and G. D. Hager, "Temporal convolutional networks for action segmentation and detection," in Proc. IEEE Conf. Computer Vision and Pattern Recognition, 2017, pp. 156-165, doi: 10.1109/CVPR.2017.113.',
    '[16] X. Song, F. Yang, D. Wang, and K.-L. Tsui, "Combined CNN-LSTM network for state-of-charge estimation of lithium-ion batteries," IEEE Access, vol. 7, pp. 88894-88902, 2019, doi: 10.1109/ACCESS.2019.2926517.',
    '[17] M. Raissi, P. Perdikaris, and G. E. Karniadakis, "Physics-informed neural networks: A deep learning framework for solving forward and inverse problems involving nonlinear partial differential equations," Journal of Computational Physics, vol. 378, pp. 686-707, Feb. 2019, doi: 10.1016/j.jcp.2018.10.045.',
    '[18] G. E. Karniadakis et al., "Physics-informed machine learning," Nature Reviews Physics, vol. 3, pp. 422-440, Jun. 2021, doi: 10.1038/s42254-021-00314-5.',
    '[19] P. J. Kollmeyer, M. Preindl, and A. Emadi, "Lithium-ion battery dataset for state of charge estimation," Mendeley Data, v1, 2020, doi: 10.17632/cp3473x7xv.1.',
    '[20] C. Vidal, P. Malysz, P. J. Kollmeyer, and A. Emadi, "Machine learning for battery state of charge estimation," in Proc. IEEE Transportation Electrification Conference, 2019, doi: 10.1109/ITEC.2019.8790578.',
    '[21] Idaho National Laboratory, "Battery test manual for electric vehicles," INL/EXT-15-34184, Rev. 3, Jun. 2015.',
    '[22] D. Runje and S. M. Shankaranarayana, "Constrained Monotonic Neural Networks," arXiv:2205.11775, revised May 2023, doi: 10.48550/arXiv.2205.11775.',
    '[23] International Organization for Standardization, "ISO/PAS 8800:2024, Road vehicles - Safety and artificial intelligence," Publicly Available Specification, Edition 1, Dec. 2024.',
    '[24] H. Aboueidah and A. Altahhan, "A Comparison of Baseline Models and a Transformer Network for SOC Prediction in Lithium-Ion Batteries," arXiv:2410.17049, Oct. 2024.',
]


def build_doc():
    doc = Document(TEMPLATE)
    clear_body(doc)
    for sec in doc.sections:
        set_page(sec)

    # Front matter: one-column, template-compatible.
    p = add_para(doc, TITLE, align=WD_ALIGN_PARAGRAPH.LEFT, font="Helvetica", size=20, bold=True)
    p.paragraph_format.space_after = Pt(6)
    add_para(doc, AUTHOR, font="Helvetica", size=9)
    add_para(doc, AFFILIATION, font="Helvetica", size=8)
    add_para(doc, "Received: DD MM YY, Revised: DD MM YY, Accepted: DD MM YY", font="Helvetica", size=8)
    add_para(doc, f"Corresponding Author: {AUTHOR} (email: [ISI EMAIL])", font="Helvetica", size=8)

    abstract = (
        "ABSTRAK - Estimasi state of charge (SOC) merupakan fungsi kunci battery management system karena SOC memengaruhi proteksi, pembatasan daya, dan interpretasi kondisi baterai Li-Ion. "
        "Pada suhu ekstrem, hubungan tegangan, arus, temperatur, dan SOC menjadi kurang teramati; estimator neural berbasis jendela juga dapat mengalami cold-start anchor berulang dan menghasilkan lintasan yang tidak konsisten terhadap arah arus. "
        "Penelitian ini mengevaluasi estimator anchor-aware Hard-Coulomb dengan calibrated carried inference untuk menggabungkan observabilitas jendela, konsistensi tanda terhadap arus terukur, dan inferensi recursive. "
        "Evaluasi dilakukan pada pipeline v5 yang mengoreksi artifact label dan decimation, menggunakan ablation bertahap, multi-seed lima seed, pembandingan null OCV+Coulomb, serta baseline EKF/1RC kontinu. "
        "Hasil menunjukkan anchor_last mencapai RMSE windowed 9.99 ± 1.09% pada Scenario A dan 4.74 ± 0.31% pada Scenario B. "
        "Pada inferensi recursive, kalibrasi η*=2.0 menurunkan RMSE calibrated recursive HC menjadi 4.43%, termasuk 3.59% pada −20 °C, sedangkan EKF terbaik yang diuji mencapai 6.85%. "
        "Seluruh angka utama ditelusurkan dari artifact v5 tanpa pelatihan ulang pada tahap penulisan. "
        "Temuan utama adalah bahwa sistem akhir anchor_last + calibrated carried inference diperoleh melalui analisis kegagalan bertahap, bukan pemilihan model arbitrer. "
        "Klaimnya dibatasi pada konsistensi tanda terhadap arus terukur dan protokol v5; kebenaran fisik di bawah sensor fault serta keselamatan fungsional tidak diklaim."
    )
    add_para(doc, abstract, font="Times New Roman", size=8)
    add_para(
        doc,
        "KATA KUNCI - Estimasi SOC, Baterai Li-Ion, Hard-Coulomb, LSTM, Anchor-Aware, Recursive Inference, EKF, Suhu Ekstrem.",
        font="Times New Roman",
        size=8,
    )

    section(doc, 2)

    add_heading(doc, "I. PENDAHULUAN", 1)
    intro_paras = [
        "Estimasi SOC adalah fungsi inti BMS karena SOC tidak dapat diukur langsung, tetapi memengaruhi estimasi jarak tempuh, batas charge-discharge, strategi proteksi, dan keputusan kontrol pada kendaraan listrik [1]-[5]. Kesalahan SOC tidak hanya menaikkan RMSE, tetapi dapat mengubah lintasan keputusan operasi, terutama ketika estimator digunakan secara recursive.",
        "Suhu ekstrem memperumit estimasi SOC karena resistansi internal, polarisasi, kapasitas efektif, dan respons tegangan terminal berubah terhadap temperatur [8]-[10]. Pada suhu rendah, voltage sag di bawah beban dapat menyerupai SOC rendah, sehingga estimator yang membaca kondisi awal jendela tanpa konteks cukup dapat membuat anchor yang bias.",
        "Model neural sekuens seperti LSTM, TCN, dan CNN-LSTM menarik karena dapat mempelajari relasi nonlinear antara tegangan, arus, temperatur, dan SOC [6], [7], [13]-[16]. Namun, akurasi regresi tidak otomatis menjamin konsistensi lintasan. Model unconstrained dapat menaikkan SOC saat discharge atau menurunkan SOC saat charge meskipun RMSE rata-rata terlihat kompetitif.",
        "Protokol windowed/Seq2Point menambah masalah lain. Setiap jendela diprediksi sebagai potongan independen dan membutuhkan anchor SOC baru. Jika jendela dimulai pada kondisi cold-load, kesalahan anchor dapat berulang pada banyak jendela, sehingga evaluasi windowed tidak selalu merepresentasikan deployment recursive.",
        "EKF dan ECM merupakan baseline klasik yang relevan karena keduanya memodelkan SOC sebagai state recursive dengan koreksi tegangan. Akan tetapi, koreksi tegangan bergantung pada parameter OCV/ECM dan tingkat kepercayaan noise pengukuran. Pada suhu ekstrem, koreksi dapat mengurangi drift tetapi juga mendorong lintasan melawan arah arus terukur.",
        "Celah penelitian ini adalah evaluasi yang menggabungkan akurasi, konsistensi lintasan, observabilitas anchor, perilaku recursive inference, dan kalibrasi laju delta. Banyak studi menekankan RMSE, sedangkan studi ini menanyakan mode gagal mana yang terjadi dan intervensi apa yang menutupnya.",
    ]
    for p in intro_paras:
        add_para(doc, p)
    add_para(doc, "Kontribusi penelitian ini adalah sebagai berikut.")
    add_numbered(
        doc,
        [
            "Mengusulkan estimator SOC Hard-Coulomb anchor-aware yang secara struktural menegakkan konsistensi tanda terhadap arus terukur.",
            "Memberikan bukti ablation bertahap yang menunjukkan mengapa post-hoc clamp, baseline neural vanilla, dan Hard-Coulomb anchor-first tidak memadai.",
            "Membangun pipeline evaluasi v5 terkoreksi yang menangani artifact label dan decimation pada representasi data yang digunakan.",
            "Menunjukkan bahwa calibrated carried inference memulihkan fidelity laju delta dan memperbaiki estimasi SOC recursive.",
            "Membandingkan sistem akhir terhadap null OCV+Coulomb dan EKF/1RC kontinu untuk mengekspos frontier correction-vs-consistency pada protokol v5.",
        ],
    )

    add_heading(doc, "II. METODE", 1)
    add_heading(doc, "A. Dataset dan Koreksi v5", 2)
    for p in [
        "Dataset legacy v4 menyimpan dua sumber bias yang memengaruhi interpretasi hasil: label awal segmen dapat terkontaminasi loaded-start ohmic bias, dan first-sample decimation dapat membuat arus representatif per detik tidak selaras dengan dinamika intradetik. Kampanye v5 membangun empat varian untuk mengisolasi koreksi label dan decimation, lalu memilih v5c sebagai dataset final.",
        "Varian v5c memakai label ohmic-corrected dan mean-per-second decimation. Koreksi ini mengurangi konflik routing pada representasi per-detik yang digunakan dan mempertahankan jumlah segmen serta split jendela agar perbandingan tetap auditable. Koreksi ohmic adalah batas bawah karena komponen polarisasi tidak dimodelkan penuh.",
    ]:
        add_para(doc, p)

    section(doc, 1)
    tables = load_tables()
    add_dataframe_table(doc, "Tabel I. Varian dataset v4-v5 dan koreksi label/decimation.", table_dataset(tables["table01_dataset_variants"]), [0.55, 0.45, 1.05, 1.05, 0.55, 0.50, 0.55, 0.70, 0.45], 6.2)
    section(doc, 2)

    add_heading(doc, "B. Formulasi SOC dan Coulomb Counting", 2)
    add_para(doc, "Konvensi arus mengikuti pipeline eksperimen: discharge bernilai negatif dan charge bernilai positif. SOC dinyatakan dalam rentang 0 sampai 1. Untuk arus terukur I_k, interval sampling Δt, kapasitas efektif Q_eff(T), dan faktor laju η, pembaruan Coulomb counting ditulis sebagai berikut.")
    add_equation(doc, "SOC_t = SOC_0 − (η Δt / (3600 Q_eff(T))) Σ_{k=1}^{t} I_k", 1)
    add_para(doc, "SOC_t adalah SOC pada timestep t, SOC_0 adalah anchor awal, dan Q_eff(T) menyatakan kapasitas efektif pada temperatur T. Persamaan ini menjadi dasar interpretasi arah dan laju perubahan SOC.")

    add_heading(doc, "C. Hard-Coulomb Constraint", 2)
    add_para(doc, "Hard-Coulomb membatasi tanda delta SOC dengan aturan arus terukur. Mask discharge dan charge didefinisikan oleh ambang arus I_ε.")
    add_equation(doc, "m_t^dis = 1(I_t < −I_ε),     m_t^chg = 1(I_t > I_ε)", 2)
    add_equation(doc, "Δŝ_t = { −|Δŝ_t| jika I_t < −I_ε; +|Δŝ_t| jika I_t > I_ε; 0 atau drift terbatas jika |I_t| ≤ I_ε }", 3)
    add_equation(doc, "ŝ_t = ŝ_0 + Σ_{k=1}^{t} Δŝ_k", 4)
    add_para(doc, "Dengan konstruksi ini, PVR 0.00% adalah sifat struktural terhadap aturan arus terukur, bukan capaian empiris. Preconditions-nya eksplisit: arus terukur diasumsikan reliabel, sensor fault tidak dicakup, magnitude delta tidak sepenuhnya dijamin tepat, dan kalibrasi η tetap diperlukan untuk rate fidelity.")

    add_heading(doc, "D. Anchor-Aware Model", 2)
    add_para(doc, "Hard-Coulomb awal memakai anchor-first, yaitu anchor SOC diprediksi dari hidden state awal. Pada kondisi cold-load, hidden state awal dapat didominasi tegangan terpolarisasi sehingga anchor tidak observabel dengan baik. Anchor-last membaca representasi h_T setelah encoder melihat seluruh jendela kausal, lalu anchor diremap ke interval feasible dari akumulasi delta.")
    add_equation(doc, "ŝ_0 = f_θ(h_T) = f_θ(Encoder(x_1:T))", 5)
    add_para(doc, "Desain ini tidak menghilangkan seluruh ambiguitas suhu dingin, tetapi memberi anchor informasi window-level yang tidak tersedia pada anchor-first.")

    add_heading(doc, "E. Calibrated Carried Inference", 2)
    add_para(doc, "Pada independent windowing, setiap jendela memulai estimasi dari anchor baru. Pada carried inference, SOC hasil jendela sebelumnya menjadi state awal jendela berikutnya sehingga artifact cold-start tidak berulang. Kebijakan load-gated melakukan re-anchor hanya ketika kondisi rest atau low-load membuat anchor tegangan lebih dapat dipercaya.")
    add_equation(doc, "ŝ_t^rec = ŝ_{t−1}^rec + η* Δŝ_t", 6)
    add_equation(doc, "r_Δ = mean(|Δŝ|) / mean(|Δs|)", 7)
    add_para(doc, "Pada bukti v5 yang tersedia, η*=2.0 dipilih karena memulihkan r_Δ dari 0.751 menjadi sekitar 1.002. Dengan carried inference, konfigurasi ini menghasilkan RMSE 4.43%. Bukti ini berasal dari Scenario A seed 42 dan harus diperlakukan sebagai single-checkpoint/single-scenario sampai re-derivation validation-chain dan konfirmasi multi-seed selesai. Inilah konfigurasi akhir anchor_last + calibrated carried inference.")

    add_heading(doc, "F. Baseline Methods", 2)
    add_para(doc, "Baseline yang diuji meliputi null OCV+Coulomb, Vanilla LSTM/TCN, post-hoc clamp, Hard-Coulomb anchor-first, HC-TCN, anchor_last, anchor_pooled, dan EKF kontinu berbasis OCV-Rint serta 1RC-ECM. EKF memakai parameter literature-like, bukan identifikasi khusus sel ini; karena itu EKF dipakai sebagai pembanding klasik yang transparan, bukan batas performa mutlak.")

    section(doc, 1)
    add_figure(doc, FIG_DIR / "fig01_final_architecture.png", "Gambar 1. Sistem akhir memakai anchor_last untuk observabilitas jendela dan carried inference terkalibrasi untuk menjaga kontinuitas recursive.", 7.0, full_width=True)
    section(doc, 2)

    add_heading(doc, "III. SETUP EKSPERIMEN", 1)
    for p in [
        "Scenario A mengevaluasi temperature-OOD, sedangkan Scenario B memakai chronological in-distribution split. Model windowed dilatih dan dievaluasi pada lima seed, yaitu 1 sampai 5, sehingga stabilitas mean dan standar deviasi dapat dibaca secara langsung.",
        "Eksperimen recursive dan η calibration memakai Scenario A seed 42 sesuai artifact v5 yang tersedia. Karena itu, angka recursive, load-gated, η*=2.0, dan EKF ditulis sebagai evidence deployment-style single-checkpoint, bukan klaim multi-seed.",
        "Metrik utama adalah RMSE, MAE, MaxE, PVR discharge, dan delta ratio. PVR dipakai untuk memeriksa apakah prediksi delta SOC melawan arah arus discharge terukur; untuk Hard-Coulomb, PVR 0.00% mengikuti konstruksi routing.",
    ]:
        add_para(doc, p)
    add_equation(doc, "RMSE = sqrt((1/N) Σ_{i=1}^{N} (s_i − ŝ_i)^2)", 8)
    add_equation(doc, "MAE = (1/N) Σ_{i=1}^{N} |s_i − ŝ_i|", 9)
    add_equation(doc, "MaxE = max_i |s_i − ŝ_i|", 10)
    add_equation(doc, "PVR_dis = Σ_t 1(I_t < −I_ε) 1(Δŝ_t > ε) / Σ_t 1(I_t < −I_ε)", 11)
    add_para(doc, "Delta ratio dihitung sesuai (7). Nilai r_Δ di bawah 1 berarti model mengecilkan magnitude pergerakan SOC, sedangkan nilai mendekati 1 menunjukkan rate fidelity yang lebih baik.")

    add_heading(doc, "IV. HASIL DAN PEMBAHASAN", 1)
    add_heading(doc, "A. Bagaimana Kegagalan Awal Mengarah ke Metode Akhir?", 2)
    add_para(doc, "Hasil v5 menunjukkan bahwa sistem akhir tidak dipilih secara arbitrer. Setiap komponen muncul untuk menutup mode gagal yang terukur: cold-start pada windowing, lintasan vanilla yang physics-blind, kelemahan penalty constraint, collapse pada post-hoc clamp, bottleneck anchor-first, dan underestimation pada delta path.")
    section(doc, 1)
    add_figure(doc, FIG_DIR / "fig02_research_evolution_flowchart.png", "Gambar 2. Alur evolusi penelitian menunjukkan bahwa anchor_last + calibrated carried inference adalah hasil analisis kegagalan bertahap.", 7.0, full_width=True)
    add_dataframe_table(doc, "Tabel II. Mode gagal, bukti, diagnosis, perbaikan, dan peran pada sistem akhir.", table_failure(tables["table02_failure_mode_rationale"]), [1.15, 1.55, 1.35, 1.25, 1.35], 6.4)
    section(doc, 2)

    add_heading(doc, "B. Apakah Koreksi Dataset v5 Mengubah Kesimpulan?", 2)
    add_para(doc, "Koreksi v5 mengubah magnitude error dan menghapus sebagian narasi MaxE katastrofik dari v4, tetapi tidak membatalkan kontribusi utama. Artifact label legacy membuat beberapa segmen dimulai di bawah beban, sedangkan mean-per-second decimation mengurangi konflik routing pada representasi yang digunakan. Dengan v5c, perbandingan model menjadi lebih adil dan tetap menunjukkan keuntungan anchor_last.")
    add_figure(doc, FIG_DIR / "fig03_dataset_v5_correction.png", "Gambar 3. Koreksi dataset v5 mengubah besaran label shift relatif v4 sambil mempertahankan split evaluasi yang dapat diaudit.", 3.25)

    add_heading(doc, "C. Apakah Constraint Perlu Dilatih End-to-End?", 2)
    add_para(doc, "Post-hoc clamp diuji sebagai alternatif sederhana untuk memaksa konsistensi setelah prediksi unconstrained. Hasilnya justru collapse, dengan RMSE 25.03% pada Scenario A dan 30.07% pada Scenario B. Ini menunjukkan bahwa filtering keluaran tidak memperbaiki representasi lintasan yang dipelajari. Constraint perlu menjadi bagian dari jalur training dan rekonstruksi SOC, bukan hanya koreksi akhir.")
    section(doc, 1)
    add_dataframe_table(doc, "Tabel III. Perbandingan model utama pada protokol v5.", table_model(tables["table03_main_model_comparison"]), [1.7, 1.1, 1.1, 2.6], 7.0)
    section(doc, 2)

    add_heading(doc, "D. Apakah Anchor-Last Stabil?", 2)
    add_para(doc, "Bukti multi-seed menunjukkan anchor_last sebagai konfigurasi windowed terbaik. Pada Scenario A, anchor_last mencapai 9.99 ± 1.09% RMSE; pada Scenario B, 4.74 ± 0.31%. Sebaliknya, Hard-Coulomb original anchor-first gagal sistematis pada Scenario B dengan 10.63 ± 0.60% pada 5/5 seed. Ini mendukung diagnosis bahwa bottleneck utama bukan hanya constraint, melainkan observabilitas anchor.")
    section(doc, 1)
    add_figure(doc, FIG_DIR / "fig04_multiseed_model_comparison.png", "Gambar 4. Perbandingan multi-seed menunjukkan anchor_last sebagai model windowed terbaik dan lebih stabil pada Scenario B.", 7.0, full_width=True)
    add_dataframe_table(doc, "Tabel IV. Ringkasan stabilitas multi-seed untuk model neural utama.", table_multiseed(tables["table04_multiseed_summary"]), [1.45, 0.85, 1.05, 1.05, 1.0, 1.0], 6.5)
    section(doc, 2)

    add_heading(doc, "E. Apakah Recursive Inference Mengurangi Cold-Start Artifact?", 2)
    add_para(doc, "Recursive inference mengurangi pengulangan cold-start karena state SOC dibawa antarjendela. Policy load-gated mencapai 8.41% RMSE dan memperbaiki kondisi dingin dibanding windowed independent, tetapi masih mempertahankan trade-off terhadap kondisi panas. Ini menunjukkan bahwa kontinuitas state penting, namun belum cukup tanpa kalibrasi laju delta.")
    add_figure(doc, FIG_DIR / "fig05_recursive_policy_comparison.png", "Gambar 5. Kebijakan recursive menurunkan artifact cold-start, dengan load-gated sebagai kompromi terbaik sebelum kalibrasi η.", 3.25)

    add_heading(doc, "F. Apakah Eta Calibration Memperbaiki Delta-Path?", 2)
    add_para(doc, "Hard-Coulomb awal mendapatkan arah delta dengan benar, tetapi rate path-nya terlalu kecil. Delta ratio 0.751 berarti magnitude gerak SOC diperkecil relatif terhadap label. Kalibrasi η*=2.0 memulihkan r_Δ menjadi sekitar 1.002 dan menurunkan RMSE recursive menjadi 4.43%, termasuk 3.59% pada −20 °C. Retraining tidak otomatis menyelesaikan masalah ini karena head model dapat mengompensasi envelope selama training.")
    add_figure(doc, FIG_DIR / "fig06_eta_calibration_delta_ratio.png", "Gambar 6. Kalibrasi η memulihkan delta ratio mendekati 1 dan menurunkan RMSE recursive pada protokol v5.", 3.25)
    section(doc, 1)
    add_dataframe_table(doc, "Tabel V. Ablation recursive policy dan η calibration.", table_recursive(tables["table05_recursive_eta_ablation"]), [1.45, 0.65, 0.65, 0.65, 0.65, 0.65, 0.80, 1.35], 6.4)
    section(doc, 2)

    add_heading(doc, "G. Bagaimana Dibandingkan dengan EKF?", 2)
    add_para(doc, "EKF adalah pembanding recursive yang wajar karena menggabungkan prediksi Coulomb dan koreksi tegangan. Pada protokol v5, EKF 1RC terbaik yang diuji mencapai 6.85% RMSE, sedangkan calibrated recursive HC mencapai 4.43%. Namun klaim ini terbatas pada parameterisasi EKF yang diuji; parameter EKF bersifat literature-like dan belum diidentifikasi khusus untuk sel ini. Perbedaan pentingnya adalah EKF memiliki PVR nonzero sekitar 5-40%, sedangkan Hard-Coulomb menjaga PVR 0.00% by construction terhadap arus terukur.")
    add_figure(doc, FIG_DIR / "fig07_ekf_vs_calibrated_hc.png", "Gambar 7. Calibrated recursive HC memiliki RMSE lebih rendah daripada EKF terbaik yang diuji dan menjaga PVR struktural 0.00% terhadap arus terukur.", 3.25)
    section(doc, 1)
    add_dataframe_table(doc, "Tabel VI. Perbandingan EKF dan sistem akhir.", table_ekf(tables["table06_ekf_comparison"]), [1.7, 0.7, 0.7, 0.7, 2.6], 7.0)
    section(doc, 2)

    add_heading(doc, "H. Apa Batasan Klaim?", 2)
    add_para(doc, "Klaim utama harus dibaca secara konservatif. Pertama, PVR 0.00% adalah konsekuensi struktural dari routing Hard-Coulomb terhadap arus terukur, bukan bukti empiris bahwa semua fisika baterai terpenuhi. Kedua, measured-current consistency tidak menyelesaikan sensor fault. Ketiga, η*=2.0 masih memerlukan re-derivation pada validation chain serta konfirmasi Scenario B dan multi-seed. Keempat, EKF yang diuji belum memakai identifikasi parameter cell-specific. Kelima, deployment edge atau keselamatan fungsional tidak diklaim.")

    add_heading(doc, "V. KESIMPULAN", 1)
    for p in [
        "Penelitian ini merekonstruksi estimator SOC Li-Ion suhu ekstrem melalui analisis kegagalan bertahap. Sistem akhir, anchor_last + calibrated carried inference, dipilih karena anchor_last memperbaiki observabilitas jendela, Hard-Coulomb menegakkan konsistensi tanda terhadap arus terukur, dan carried inference terkalibrasi mengurangi cold-start berulang sekaligus memulihkan laju delta.",
        "Pada protokol v5, anchor_last mencapai 9.99 ± 1.09% RMSE pada Scenario A dan 4.74 ± 0.31% pada Scenario B. Pada evaluasi recursive Scenario A seed 42, η*=2.0 menghasilkan RMSE 4.43% dan 3.59% pada −20 °C. Dibandingkan baseline EKF/1RC terbaik yang diuji, sistem akhir memiliki RMSE lebih rendah dan menjaga PVR 0.00% sebagai sifat struktural terhadap arus terukur.",
        "Pekerjaan lanjut mencakup derivasi η* pada validation chain, konfirmasi multi-seed dan Scenario B untuk calibrated carried inference, diagnosis sensor fault, identifikasi parameter ECM yang cell-specific, deployment terkuantisasi dengan akumulator sign-preserving, serta evaluasi pada dataset baterai yang lebih luas.",
    ]:
        add_para(doc, p)

    add_heading(doc, "KONFLIK KEPENTINGAN", 1)
    add_para(doc, "Penulis menyatakan tidak terdapat konflik kepentingan.")

    add_heading(doc, "KONTRIBUSI PENULIS", 1)
    add_para(doc, "Abyan Hisyam Al'ammar: konseptualisasi, metodologi, perangkat lunak, validasi, analisis formal, investigasi, kurasi data, visualisasi, penulisan draf awal, dan penyuntingan naskah. [ISI NAMA DOSEN PEMBIMBING]: supervisi, validasi metodologi, telaah naskah, dan arahan penelitian.")

    add_heading(doc, "UCAPAN TERIMA KASIH", 1)
    add_para(doc, "Penulis mengucapkan terima kasih kepada [ISI NAMA DOSEN/PROGRAM STUDI/LAB] atas arahan dan masukan selama proses penelitian.")

    add_heading(doc, "REFERENSI", 1)
    for ref in REFERENCES:
        p = add_para(doc, ref, font="Times New Roman", size=8)
        p.paragraph_format.left_indent = Mm(5)
        p.paragraph_format.first_line_indent = Mm(-5)

    add_heading(doc, "LAMPIRAN", 1)
    add_para(doc, "Lampiran ini merangkum bukti reproduksibilitas yang tidak dimuat penuh pada teks utama. Final ablation matrix berisi 62 baris; claims register v2 berisi 17 klaim; readiness gate Phase 10 melaporkan 52/52 PASS. Notebook pendukung berada pada notebooks/ablation_studies_v5_final/12-19.")
    add_bullets(
        doc,
        [
            "P0-P10: freeze legacy, koreksi dataset, retraining v5, multi-seed, recursive policy, η calibration, EKF baseline, final ablation, claims register, final report, readiness gate.",
            "Artifact utama: results/v5/final_ablation_matrix.csv, results/v5/multiseed/multiseed_summary.csv, results/v5/recursive_inference/, results/v5/delta_calibration/, dan results/v5/ekf_ecm/.",
            "Batasan terbuka: η* validation-chain, konfirmasi multi-seed Scenario B, sensor fault, dan identifikasi parameter EKF/ECM cell-specific.",
        ],
    )
    section(doc, 1)
    add_dataframe_table(doc, "Tabel VII. Ringkasan klaim dan batasan dari claims register v2.", table_claims(tables["table07_claim_summary"]), [0.35, 0.75, 2.35, 3.10], 6.2)

    # Basic document-level metadata.
    doc.core_properties.title = TITLE
    doc.core_properties.author = AUTHOR
    doc.core_properties.subject = "SOC estimation, Hard-Coulomb, battery management system"
    doc.save(OUT)


def validate_docx_text():
    doc = Document(OUT)
    text = "\n".join(p.text for p in doc.paragraphs)
    def allowed_bracket(x: str) -> bool:
        if x in {"[ISI EMAIL]", "[ISI NAMA DOSEN PEMBIMBING]", "[ISI NAMA DOSEN/PROGRAM STUDI/LAB]"}:
            return True
        return bool(re.fullmatch(r"\[\d+(?:\]-\[\d+|\])?", x))

    checks = {
        "double_heading": not bool(re.search(r"\b([IVX]+\.|[A-Z]\.)\s+\1", text)),
        "no_ref_needed": "[REF NEEDED]" not in text,
        "no_backticks": "`" not in text,
        "final_system": "anchor_last + calibrated carried inference" in text,
        "no_duplicate_reference": not bool(re.search(r"\[\d+\]\s+\[\d+\]", text)),
        "no_unsafe_claims": not any(
            phrase in text.lower()
            for phrase in [
                "iso 26262 compliant",
                "functional safety guaranteed",
                "universally outperforms",
                "physically guaranteed",
                "edge-ready",
                "fail-safe",
            ]
        ),
        "intentional_placeholders_only": all(allowed_bracket(x) for x in re.findall(r"\[[^\]]+\]", text)),
    }
    return checks


def write_traceability(checks):
    fig_manifest = json.loads((ASSET_DIR / "figure_manifest.json").read_text(encoding="utf-8"))
    table_manifest = json.loads((ASSET_DIR / "table_manifest.json").read_text(encoding="utf-8"))
    lines = [
        "# Traceability v3",
        "",
        "## Output",
        f"- DOCX: `{OUT.relative_to(ROOT)}`",
        "",
        "## Figures used",
    ]
    for item in fig_manifest:
        lines.append(f"- `{item['asset_filename']}` -> {item['scientific_message']} Placement: {item['recommended_manuscript_placement']}. Caveat: {item['caveat']}")
    lines += ["", "## Tables used"]
    for item in table_manifest:
        lines.append(f"- `{item['asset_filename']}` -> {item['scientific_message']} Caveat: {item['caveat']}")
    lines += [
        "",
        "## Result/report files used",
        "- `outputs/manuscript_assets/figure_manifest.json`",
        "- `outputs/manuscript_assets/table_manifest.json`",
        "- `outputs/manuscript_assets/asset_export_report.md`",
        "- `reports/v5_campaign/phase9_final_v5_report.md`",
        "- `reports/v5_campaign/claims_register_v2.md` and `.json`",
        "- `reports/v5_campaign/phase10_manuscript_readiness_gate.md`",
        "- `drafts/JNTETI_SOC_Hard_Coulomb_v5_Draft_ID_v2.docx` as narrative source only",
        "",
        "## Numbers inserted",
        "- anchor_last Scenario A: 9.99 ± 1.09% RMSE",
        "- anchor_last Scenario B: 4.74 ± 0.31% RMSE",
        "- Original HC Scenario B failure: 10.63 ± 0.60% RMSE, 5/5 seeds",
        "- Post-hoc clamp: 25.03% / 30.07% RMSE",
        "- load_gated recursive: 8.41% RMSE",
        "- η*=2.0: delta ratio about 1.002; calibrated recursive RMSE 4.43%; −20 °C RMSE 3.59%",
        "- Best EKF 1RC R=1e-2: 6.85% RMSE; PVR 5.48%",
        "- Final ablation matrix: 62 rows; claims register: 17 claims; readiness gate: 52/52 PASS",
        "",
        "## Manual assumptions",
        "- References preserved from v2 and not re-verified online in this DOCX rebuild step.",
        "- Main manuscript tables are simplified views derived from CSV exports to keep JNTETI layout readable; full CSV files remain authoritative.",
        "- Recursive, η calibration, and EKF comparisons are labeled as Scenario A seed 42 single-checkpoint evidence.",
        "",
        "## Unresolved placeholders intentionally retained",
        "- `[ISI EMAIL]`",
        "- `[ISI NAMA DOSEN PEMBIMBING]`",
        "- `[ISI NAMA DOSEN/PROGRAM STUDI/LAB]`",
        "",
        "## Missing artifacts",
        "- None for required v3 figures/tables.",
        "- Optional PDF preview was not created because no local LibreOffice/Word render command is available.",
        "",
        "## Validation checks",
    ]
    for k, v in checks.items():
        lines.append(f"- {k}: {'PASS' if v else 'FAIL'}")
    TRACE.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_revision_notes(checks):
    lines = [
        "# Revision Notes v3",
        "",
        "## Fixed",
        "- Rebuilt manuscript from `Template-JNTETI-2025-ENG (3).docx` instead of patching v2 layout.",
        "- Removed double heading numbering by using manual heading text only.",
        "- Rebuilt all manuscript tables from `outputs/tables/*.csv`; first-column labels retained.",
        "- Inserted only clean publication figures from `outputs/figures/fig01...fig07.png`.",
        "- Replaced raw LaTeX-style equations with readable centered equation blocks numbered (1)-(11).",
        "- Reworked Results into question-driven subsections and added mandatory failure-mode rationale table.",
        "- Removed `[REF NEEDED]` and duplicate reference numbering patterns.",
        "- Kept claims conservative: PVR is structural/by construction, EKF comparison limited, and η*=2.0 caveat retained.",
        "",
        "## Remaining limitations before supervisor review",
        "- Author email, supervisor name, and acknowledgement institution/lab placeholders remain.",
        "- References were preserved from v2 and should be checked by supervisor/reference manager before submission.",
        "- η*=2.0 still needs validation-chain derivation and multi-seed/Scenario B confirmation.",
        "- EKF parameters are literature-like, not cell-identified.",
        "- Visual render QA/PDF export could not be completed locally because LibreOffice/Word command-line rendering is unavailable.",
        "",
        "## Remaining limitations before journal submission",
        "- Verify JNTETI final page layout in Microsoft Word after any manual edits.",
        "- Update received/revised/accepted dates if required by the journal workflow.",
        "- Re-check reference metadata and citation coverage.",
        "- Consider adding a cell-specific ECM/EKF identification experiment if reviewers demand a stronger classical baseline.",
        "",
        "## Suitability",
        "- Suitable for supervisor review: yes.",
        "- Suitable for immediate journal submission: not yet; placeholders and reference audit remain.",
        "",
        "## Automated checks",
    ]
    for k, v in checks.items():
        lines.append(f"- {k}: {'PASS' if v else 'FAIL'}")
    NOTES.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    build_doc()
    checks = validate_docx_text()
    write_traceability(checks)
    write_revision_notes(checks)
    print(json.dumps(checks, indent=2))
    if not all(checks.values()):
        raise SystemExit(2)


if __name__ == "__main__":
    main()
